import streamlit as st
import pandas as pd
import numpy as np
from copy import deepcopy
from datetime import date
from calendar import monthrange
from pathlib import Path
from tempfile import gettempdir
import json
import ast
import csv
from joblib import load
import re  # ✅ NUEVO
from datetime import datetime
import os
import html as html_lib
from io import BytesIO
import xml.etree.ElementTree as ET
import zlib
import subprocess
import shutil
import tempfile
import importlib.util
from urllib.parse import urlparse, parse_qs

import gspread
from google.oauth2.service_account import Credentials
from google.oauth2.credentials import Credentials as UserCredentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import streamlit.components.v1 as components
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

# ==== Transformadores CUSTOM (deben estar antes de load) ====
from sklearn.base import BaseEstimator, TransformerMixin

class LogAndDrop(BaseEstimator, TransformerMixin):
    """
    Crea columnas log1p de ['C/A', 'AMOUNT_TOTAL'] y elimina las originales.
    Mantiene ['PRI-ULT','Ratio_PP'] tal cual.
    Salida: ['PRI-ULT','Ratio_PP','C/A_log','AMOUNT_TOTAL_log']
    """
    def __init__(self, ca_col="C/A", amt_col="AMOUNT_TOTAL"):
        self.ca_col = ca_col
        self.amt_col = amt_col
        self.out_feature_names_ = ["PRI-ULT", "Ratio_PP", "C/A_log", "AMOUNT_TOTAL_log"]

    def fit(self, X, y=None):
        return self

    def transform(self, X):
        X = X.copy()
        X["C/A_log"] = np.log1p(pd.to_numeric(X[self.ca_col], errors="coerce").astype(float))
        X["AMOUNT_TOTAL_log"] = np.log1p(pd.to_numeric(X[self.amt_col], errors="coerce").astype(float))
        X = X.drop(columns=[self.ca_col, self.amt_col])
        return X[["PRI-ULT", "Ratio_PP", "C/A_log", "AMOUNT_TOTAL_log"]]

    def get_feature_names_out(self, input_features=None):
        return np.array(self.out_feature_names_)

class Winsorizer(BaseEstimator, TransformerMixin):
    """
    Winsoriza por cuantiles (p_low, p_high) columnas numéricas.
    Aprende límites en fit y los aplica en transform.
    """
    def __init__(self, columns=None, p_low=0.005, p_high=0.005):
        self.columns = columns or []
        self.p_low = p_low
        self.p_high = p_high
        self.lows_ = {}
        self.highs_ = {}

    def fit(self, X, y=None):
        X = pd.DataFrame(X, copy=True)
        for c in self.columns:
            s = pd.to_numeric(X[c], errors="coerce")
            self.lows_[c] = s.quantile(self.p_low)
            self.highs_[c] = s.quantile(1 - self.p_high)
        return self

    def transform(self, X):
        X = pd.DataFrame(X, copy=True)
        for c in self.columns:
            lo = self.lows_[c]
            hi = self.highs_[c]
            X[c] = pd.to_numeric(X[c], errors="coerce").clip(lower=lo, upper=hi)
        return X


# ------------------ Ajustes globales ------------------
pd.set_option("mode.copy_on_write", True)
def _configure_streamlit_page():
    try:
        st.set_page_config(page_title="Calculadora de Recaudo", page_icon="💸", layout="centered")
    except Exception:
        return


_configure_streamlit_page()
st.title("💸 Calculadora de Recaudo")

import sklearn, numpy, joblib
from sklearn.impute import SimpleImputer

# 🔧 PARCHE compatibilidad modelo viejo vs sklearn nuevo
if not hasattr(SimpleImputer, "_fill_dtype"):
    SimpleImputer._fill_dtype = None
st.sidebar.caption(
    f"🧩 NumPy: {numpy.__version__}\n"
    f"🧠 scikit-learn: {sklearn.__version__}\n"
    f"💼 joblib: {joblib.__version__}"
)

st.caption(
    "1) La app carga automáticamente la base generada por el workflow (`data/cartera_asignada_filtrada`) • "
    "2) Escribe la **Referencia** y selecciona **uno o varios Id deuda** • "
    "3) Ajusta valores editables (Deuda, Apartado, Comisión, Saldo) • "
    "4) Ingresa **PAGO BANCO** y **N PaB** → se calcula **DESCUENTO** y **Comisión de éxito** • "
    "6) Revisa KPIs (PLAZO lo ingresas tú)."
)

# =================== 🔄 Reinicio manual (limpiar cache) ===================
st.sidebar.markdown("### 🔄 Control")
if st.sidebar.button("Reiniciar calculadora (limpiar cache)"):
    st.cache_data.clear()
    st.cache_resource.clear()
    st.rerun()

# ==== Rutas de artefactos generados por el notebook/Action ====
DATA_PARQUET = Path("data/cartera_asignada_filtrada.parquet")
DATA_CSV     = Path("data/cartera_asignada_filtrada.csv")
MODEL_PATH   = Path("mlp_recaudo_pipeline.joblib")
META_PATH    = Path("mlp_recaudo_meta.json")
LOG_PATH     = Path("data/logs/logs_calculadora.csv")
DOCX_TEMPLATE_PATH = Path("data/Documento Estructurados en Blanco.docx")
CLIENTES_LOOKUP_PATH = Path("data/Consulta_F_Clientes_Parte_1.csv")
GOOGLE_SHEET_ID = "1Aahltn7TSRf6ZpTpS-vPgpB89hO-r5KxpAhqKAPXziE"
GOOGLE_SHEET_TAB = "Historico Calculadora"
GOOGLE_SHEET_TAB_RESPUESTAS = "Respuestas Estr"
GOOGLE_SHEETS_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]
GOOGLE_DRIVE_UPLOAD_SCOPES = ["https://www.googleapis.com/auth/drive.file"]
DRIVE_FOLDER_CARTA_PAGARE_ID = "1nEo1iZWzFySJX_90crO9tjTTX1Cr_yVxs-xyn1C0TMu78Jt8rs2QYqVXs_wgzxEvn1AU0nMk"
DRIVE_FOLDER_PANTALLAZOS_ID = "1wTIUNP74ZD2MtVO_bOtowM-z9z0RgpxhEarfoElwQGE86kpMiPWz7qt4130YFYK6NiXZNRh1"
DRIVE_FOLDER_CONDONACION_CORREO_ID = "1CN73OI6DjyEVGLsu1m9iszPPJ4xKfTZM5aMQ-lHNwHnlgzck0VjdL3MX5RuObC_n3zs-MFNF"

GOOGLE_SHEET_HEADERS = [
    "fecha",
    "referencia",
    "ids_deuda",
    "plazo",
    "ratio_pp",
    "cuota_apartado",
    "amount_total",
    "prediccion",
]

GOOGLE_RESPUESTAS_COLS = [chr(i) for i in range(ord("A"), ord("V") + 1)]


# ========= Helpers de "versión de archivo" para invalidar cache =========

def _file_version(path: Path) -> str:
    """
    Devuelve una 'firma' del archivo basada en mtime + tamaño.
    Si el archivo no existe, devuelve 'missing'.
    """
    try:
        stat = path.stat()
        return f"{stat.st_mtime_ns}-{stat.st_size}"
    except FileNotFoundError:
        return "missing"
##################################################################################################################################################################################        
def _sum_rounded_parts(values, digits=2):
    rounded = [round(float(v), digits) for v in values]
    if rounded:
        rounded[-1] = round(sum(values) - sum(rounded[:-1]), digits)
    return rounded

def _split_total_by_weights(total: float, weights: list[float], digits: int = 2) -> list[float]:
    total = max(float(total or 0.0), 0.0)
    if not weights:
        return []
    clean = [max(float(w or 0.0), 0.0) for w in weights]
    w_sum = sum(clean)
    if w_sum <= 0:
        clean = [1.0] * len(weights)
        w_sum = float(len(weights))
    raw = [total * (w / w_sum) for w in clean]
    return _sum_rounded_parts(raw, digits=digits)


def _split_integer_equitable(total: int, n_parts: int) -> list[int]:
    total = max(int(total or 0), 0)
    n_parts = max(int(n_parts or 0), 0)
    if n_parts == 0:
        return []
    base = total // n_parts
    rem = total % n_parts
    return [base + (1 if i < rem else 0) for i in range(n_parts)]


def _last_day_of_month(base_date: date, months_ahead: int) -> date:
    shifted = pd.Timestamp(base_date) + pd.DateOffset(months=months_ahead)
    return date(int(shifted.year), int(shifted.month), monthrange(int(shifted.year), int(shifted.month))[1])


def _day_of_month(base_date: date, months_ahead: int, day: int | None) -> date:
    if day is None:
        return _last_day_of_month(base_date, months_ahead)
    shifted = pd.Timestamp(base_date) + pd.DateOffset(months=months_ahead)
    year = int(shifted.year)
    month = int(shifted.month)
    safe_day = min(max(int(day), 1), monthrange(year, month)[1])
    return date(year, month, safe_day)


def _month_offset(base_date: date, target_date: date) -> int:
    return ((target_date.year - base_date.year) * 12) + (target_date.month - base_date.month)


def _rebalance_group_amounts(df_group: pd.DataFrame, total_objetivo: float) -> pd.DataFrame:
    df_group = df_group.sort_values("orden").copy()
    total_objetivo = max(float(total_objetivo or 0.0), 0.0)
    override_mask = df_group["cantidad_editada"].fillna(False).astype(bool)
    total_editado = min(max(float(df_group.loc[override_mask, "Cantidad"].sum()), 0.0), total_objetivo)
    restantes = df_group.index[~override_mask].tolist()
    restante_disponible = max(0.0, total_objetivo - total_editado)

    if restantes:
        partes = _sum_rounded_parts([restante_disponible / len(restantes)] * len(restantes))
        for idx, valor in zip(restantes, partes):
            df_group.at[idx, "Cantidad"] = valor
    elif override_mask.any():
        ultimo_idx = df_group.index[-1]
        otros = float(df_group.iloc[:-1]["Cantidad"].sum()) if len(df_group) > 1 else 0.0
        df_group.at[ultimo_idx, "Cantidad"] = round(max(total_objetivo - otros, 0.0), 2)

    return df_group


def _parse_amount_input(value, *, max_decimals: int = 2) -> float:
    if value is None:
        return 0.0
    if isinstance(value, (int, float, np.integer, np.floating)):
        return float(value)
    text = str(value).strip()
    if not text:
        return 0.0
    sign = -1 if text.startswith("-") else 1
    cleaned = re.sub(r"[^\d,.\-]", "", text)
    if cleaned in {"", "-", ",", ".", "-,", "-."}:
        return 0.0
    cleaned = cleaned.replace("-", "")
    if "," in cleaned and "." in cleaned:
        if cleaned.rfind(",") > cleaned.rfind("."):
            cleaned = cleaned.replace(".", "").replace(",", ".")
        else:
            cleaned = cleaned.replace(",", "")
    elif "," in cleaned:
        cleaned = cleaned.replace(".", "").replace(",", ".")
    else:
        parts = cleaned.split(".")
        if len(parts) > 2:
            cleaned = "".join(parts)
        elif len(parts) == 2 and len(parts[-1]) == 3 and parts[0].isdigit():
            cleaned = cleaned.replace(".", "")

    try:
        parsed = float(cleaned)
    except ValueError:
        return 0.0

    factor = 10 ** max(0, int(max_decimals))
    parsed = np.trunc(parsed * factor) / factor
    return sign * parsed


def _format_number_es(value, *, max_decimals: int = 2, trim_trailing: bool = True) -> str:
    amount = float(value or 0.0)
    max_decimals = max(0, int(max_decimals))
    formatted = f"{amount:,.{max_decimals}f}".replace(",", "_").replace(".", ",").replace("_", ".")
    if trim_trailing and "," in formatted:
        formatted = formatted.rstrip("0").rstrip(",")
    return formatted


def _format_currency0(value, *, decimals: int = 2, trim_trailing: bool = True) -> str:
    amount = float(value or 0.0)
    formatted = _format_number_es(amount, max_decimals=decimals, trim_trailing=trim_trailing)
    return f"$ {formatted}"


#############################################################################################################################################################################
#############################################################################################################################################################################
def _format_currency_cop(value) -> str:
    amount = float(value or 0.0)
    formatted = _format_number_es(amount, max_decimals=2, trim_trailing=True)
    return f"${formatted} COP"


def _number_to_words_es(value: int) -> str:
    units = ["cero", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve"]
    teens = {
        10: "diez", 11: "once", 12: "doce", 13: "trece", 14: "catorce", 15: "quince",
        16: "dieciseis", 17: "diecisiete", 18: "dieciocho", 19: "diecinueve",
    }
    tens = {20: "veinte", 30: "treinta", 40: "cuarenta", 50: "cincuenta", 60: "sesenta", 70: "setenta", 80: "ochenta", 90: "noventa"}
    hundreds = {100: "cien", 200: "doscientos", 300: "trescientos", 400: "cuatrocientos", 500: "quinientos", 600: "seiscientos", 700: "setecientos", 800: "ochocientos", 900: "novecientos"}

    def convert(n: int) -> str:
        if n < 10:
            return units[n]
        if n < 20:
            return teens[n]
        if n < 30:
            return "veinte" if n == 20 else f"veinti{convert(n - 20)}"
        if n < 100:
            return tens[(n // 10) * 10] if n % 10 == 0 else f"{tens[(n // 10) * 10]} y {convert(n % 10)}"
        if n == 100:
            return "cien"
        if n < 1000:
            base = hundreds.get((n // 100) * 100, "ciento")
            if (n // 100) * 100 == 100:
                base = "ciento"
            return base if n % 100 == 0 else f"{base} {convert(n % 100)}"
        if n < 1_000_000:
            thousands = n // 1000
            remainder = n % 1000
            prefix = "mil" if thousands == 1 else f"{convert(thousands)} mil"
            return prefix if remainder == 0 else f"{prefix} {convert(remainder)}"
        if n < 1_000_000_000:
            millions = n // 1_000_000
            remainder = n % 1_000_000
            prefix = "un millon" if millions == 1 else f"{convert(millions)} millones"
            return prefix if remainder == 0 else f"{prefix} {convert(remainder)}"
        return str(n)

    return convert(max(int(value), 0))


def _format_currency_cop_words(value) -> str:
    return _number_to_words_es(int(round(float(value or 0.0)))).upper()


def _format_date_ddmmyyyy(value) -> str:
    if pd.isna(value):
        return ""
    return pd.to_datetime(value).strftime("%d/%m/%Y")


def _table_cell_text(value) -> str:
    if value is None or pd.isna(value):
        return ""
    if isinstance(value, (int, float, np.integer, np.floating)):
        return _format_currency0(value)
    return str(value)


def _format_month_name_es(value) -> str:
    month_names = {
        1: "enero",
        2: "febrero",
        3: "marzo",
        4: "abril",
        5: "mayo",
        6: "junio",
        7: "julio",
        8: "agosto",
        9: "septiembre",
        10: "octubre",
        11: "noviembre",
        12: "diciembre",
    }
    dt_value = pd.to_datetime(value)
    return month_names.get(int(dt_value.month), "")


def _normalize_template_value(value) -> str:
    if value is None or (isinstance(value, float) and np.isnan(value)):
        return ""
    return str(value)



def _render_template_text(text: str, context: dict[str, str]) -> str:
    rendered = str(text)
    for key, value in context.items():
        rendered = rendered.replace(f"{{{key}}}", _normalize_template_value(value))
    return rendered


def _join_unique_values(values, separator=" - ") -> str:
    unique_values = []
    for value in values:
        text_value = str(value).strip()
        if not text_value or text_value.lower() == "nan":
            continue
        if text_value not in unique_values:
            unique_values.append(text_value)
    return separator.join(unique_values)


def _smart_title_case(value: str) -> str:
    words = []
    for word in str(value or "").split():
        parts = [part.capitalize() for part in word.split("-")]
        words.append("-".join(parts))
    return " ".join(words)


def _paragraph_has_column_break(paragraph) -> bool:
    return 'w:br w:type="column"' in paragraph._p.xml


def _build_document_context(
    referencia,
    bancos,
    pago_banco,
    comision_total,
    nombre_cliente="",
    numero_producto="",
    vehiculo="",
    cedula_cliente="",
    correo_cliente="",
    telefono_cliente="",
    ciudad_cliente="",
    direccion_cliente="",
    suma_comisiones_total=None,
) -> dict[str, str]:
    today = date.today()
    bancos_unicos = _join_unique_values(bancos)
    comision_total_text = _format_currency_cop(comision_total)
    suma_comisiones_value = float(suma_comisiones_total if suma_comisiones_total is not None else comision_total)
    suma_comisiones_text = _format_currency_cop(suma_comisiones_value)

    return {
        "referencia": str(referencia or ""),
        "dia_firma": str(today.day),
        "mes_firma": _format_month_name_es(today),
        "anio_firma": str(today.year),
        "entidad_financiera": bancos_unicos,
        "pago_banco": _format_currency_cop(pago_banco),
        "comision_total": comision_total_text,
        "nombre_cliente": _smart_title_case(nombre_cliente),
        "numero_producto": str(numero_producto or ""),
        "vehiculo": _smart_title_case(vehiculo),
        "cedula_cliente": str(cedula_cliente or ""),
        "correo_cliente": str(correo_cliente or ""),
        "telefono_cliente": str(telefono_cliente or ""),
        "ciudad_cliente": str(ciudad_cliente or ""),
        "direccion_cliente": str(direccion_cliente or ""),
        "suma_comisiones": suma_comisiones_text,
        "Suma_comisiones": suma_comisiones_text,
        "suma comisiones": suma_comisiones_text,
        "suma_comisiones_letras": _format_currency_cop_words(suma_comisiones_value),
    }


def _replace_paragraph_text_preserving_style(paragraph, new_text: str):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    template_run = next((run for run in paragraph.runs if run.text), paragraph.runs[0])

    if _paragraph_has_column_break(paragraph):
        for run in paragraph.runs:
            if run.text:
                run.text = ""
        new_run = paragraph.add_run(new_text)
        new_run.bold = template_run.bold
        new_run.italic = template_run.italic
        new_run.underline = template_run.underline
        if template_run.font is not None:
            new_run.font.name = template_run.font.name
            new_run.font.size = template_run.font.size
        return

    first_run = paragraph.runs[0]
    first_run.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _apply_context_to_paragraph(paragraph, context: dict[str, str]):
    full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    rendered_text = _render_template_text(full_text, context)
    if rendered_text != full_text:
        _replace_paragraph_text_preserving_style(paragraph, rendered_text)


def _apply_context_to_table(table, context: dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _apply_context_to_paragraph(paragraph, context)
            for nested_table in cell.tables:
                _apply_context_to_table(nested_table, context)


def _apply_context_to_document(document, context: dict[str, str]):
    for paragraph in document.paragraphs:
        _apply_context_to_paragraph(paragraph, context)
    for table in document.tables:
        _apply_context_to_table(table, context)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            _apply_context_to_paragraph(paragraph, context)
        for table in section.header.tables:
            _apply_context_to_table(table, context)
        for paragraph in section.footer.paragraphs:
            _apply_context_to_paragraph(paragraph, context)
        for table in section.footer.tables:
            _apply_context_to_table(table, context)


def _set_table_cell_no_wrap(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = tc_pr.find(qn("w:noWrap"))
    if no_wrap is None:
        no_wrap = OxmlElement("w:noWrap")
        tc_pr.append(no_wrap)


def _set_cell_width(cell, width_inches: float):
    width_twips = int(width_inches * 1440)
    cell.width = width_twips
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_twips))

def _set_table_grid_widths(table, column_widths: list[float]):
    """
    Ajusta el ancho de columnas a nivel de grid de tabla (w:tblGrid),
    que es lo que LibreOffice suele respetar al exportar a PDF.
    """
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)

    # Reemplazar definición previa del grid para evitar conflictos de ancho.
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)

    for width in column_widths:
        width_twips = int(width * 1440)
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_twips))
        tbl_grid.append(grid_col)

    # También aplicar a la API de python-docx por compatibilidad adicional.
    for idx, width in enumerate(column_widths):
        if idx < len(table.columns):
            table.columns[idx].width = int(width * 1440)


def _set_table_fixed_layout(table, column_widths: list[float]):
    """
    Fuerza layout fijo de tabla y ancho total explícito.
    Esto evita que LibreOffice "recalcule" columnas al exportar a PDF.
    """
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    total_twips = str(int(sum(column_widths) * 1440))
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), total_twips)


def _apply_cronograma_table_layout(table):
    table.autofit = False
    column_widths = [0.23, 0.9, 1.2, 1.7]
    _set_table_fixed_layout(table, column_widths)
    _set_table_grid_widths(table, column_widths)
    for row in table.rows:
        for idx, width in enumerate(column_widths):
            if idx < len(row.cells):
                _set_cell_width(row.cells[idx], width)


def _apply_table_text_style(paragraph, run, *, bold=None):
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.alignment = paragraph.alignment

    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(9.5)
    if bold is not None:
        run.bold = bold


def _replace_cell_text_preserving_style(cell, text: str):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    template_run = paragraph.runs[0] if paragraph.runs else None

    for extra_paragraph in cell.paragraphs[1:]:
        extra_paragraph._element.getparent().remove(extra_paragraph._element)

    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    normalized_text = str(text).replace("\n", " ").strip()
    new_run = paragraph.add_run(normalized_text)
    _set_table_cell_no_wrap(cell)

    inherited_bold = template_run.bold if template_run is not None else None
    if template_run is not None:
        new_run.italic = template_run.italic
        new_run.underline = template_run.underline
    _apply_table_text_style(paragraph, new_run, bold=inherited_bold)


def _populate_docx_table(table, rows: list[list[str]]):
    if len(table.rows) < 2:
        raise ValueError("La plantilla Word debe tener encabezado + una fila de muestra por tabla.")

    template_row = table.rows[1]
    while len(table.rows) > 2:
        table._tbl.remove(table.rows[-1]._tr)

    first_data_row = table.rows[1]
    if not rows:
        rows = [[""] * len(first_data_row.cells)]

    for row_idx, values in enumerate(rows):
        target_row = first_data_row if row_idx == 0 else None
        if target_row is None:
            new_tr = deepcopy(template_row._tr)
            table._tbl.append(new_tr)
            target_row = table.rows[-1]

        for col_idx, value in enumerate(values):
            if col_idx < len(target_row.cells):
                _replace_cell_text_preserving_style(target_row.cells[col_idx], value)


def _remove_paragraph(paragraph: Paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _remove_graduation_section(document):
    start_markers = [
        "mi estatus de Cliente Activo pasa a ser como Cliente Graduado",
        "Cliente Graduado",
    ]
    end_markers = [
        "Ref No",
        "C.C.",
        "Atentamente",
    ]

    paragraphs = list(document.paragraphs)
    start_idx = None
    end_idx = None

    for idx, paragraph in enumerate(paragraphs):
        text = str(paragraph.text or "").strip()
        if not text:
            continue
        if start_idx is None and any(marker in text for marker in start_markers):
            start_idx = idx
        if start_idx is not None and any(marker in text for marker in end_markers):
            end_idx = idx

    if start_idx is None:
        return

    if end_idx is None:
        end_idx = start_idx
        while end_idx + 1 < len(paragraphs) and not str(paragraphs[end_idx + 1].text or "").strip():
            end_idx += 1

    for paragraph in paragraphs[start_idx : end_idx + 1]:
        _remove_paragraph(paragraph)              


def build_recaudo_docx(
    template_path: Path,
    cronograma_df: pd.DataFrame,
    plan_df: pd.DataFrame,
    template_context: dict[str, str],
    include_graduation_section: bool = False,
) -> bytes:
    if not template_path.exists():
        raise FileNotFoundError(f"No encontré la plantilla Word: {template_path}")

    document = Document(str(template_path))
    _apply_context_to_document(document, template_context)
    if not include_graduation_section:
        _remove_graduation_section(document)
    if len(document.tables) < 2:
        raise ValueError("La plantilla Word debe tener al menos dos tablas para reemplazar.")

    cronograma_export = cronograma_df[cronograma_df["Cantidad"] > 0.005][["Fecha", "Cantidad", "Concepto"]].copy()
    cronograma_rows = []
    for idx, row in cronograma_export.reset_index(drop=True).iterrows():
        cronograma_rows.append([
            str(idx + 1),
            _format_date_ddmmyyyy(row["Fecha"]),
            _format_currency_cop(row["Cantidad"]),
            str(row["Concepto"]),
        ])

    plan_export = plan_df.copy()
    plan_rows = []
    for _, row in plan_export.iterrows():
        plan_rows.append([
            "",
            _format_date_ddmmyyyy(row["Fecha Límite de Pago"]),
            _table_cell_text(row["Pago a Banco"]),
            _table_cell_text(row["Comisión de Éxito"]),
            _table_cell_text(row["Comisión Mensual"]),
            _table_cell_text(row["Apartado Requerido"]),
        ])

    _populate_docx_table(document.tables[0], cronograma_rows)
    _apply_cronograma_table_layout(document.tables[0])
    _populate_docx_table(document.tables[1], plan_rows)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()


def _element_text(elem) -> str:
    texts = []
    for node in elem.iter():
        if node.tag.endswith("}t") and node.text:
            texts.append(str(node.text))
    return "".join(texts).strip()


def _build_page_break_paragraph():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _extract_first_liquidacion_block_elements(document: Document) -> list:
    body_children = list(document.element.body.iterchildren())
    first_table_idx = next((i for i, child in enumerate(body_children) if child.tag.endswith("}tbl")), None)
    if first_table_idx is None:
        return []
    start_idx = next(
        (
            i for i, child in enumerate(body_children[: first_table_idx + 1])
            if child.tag.endswith("}p") and ("Liquidación de Deuda" in _element_text(child) or "Colombia, Bogotá" in _element_text(child))
        ),
        0,
    )
    end_idx = next(
        (
            i for i, child in enumerate(body_children[first_table_idx + 1 :], start=first_table_idx + 1)
            if child.tag.endswith("}p") and _element_text(child).startswith("Ref No")
        ),
        first_table_idx,
    )
    return [deepcopy(child) for child in body_children[start_idx : end_idx + 1]]


def insert_extra_liquidacion_blocks(base_docx: bytes, extra_blocks: list[list]) -> bytes:
    document = Document(BytesIO(base_docx))
    body = document.element.body
    children = list(body.iterchildren())
    first_table_idx = next((i for i, child in enumerate(children) if child.tag.endswith("}tbl")), None)
    if first_table_idx is None:
        return base_docx
    insert_idx = next(
        (
            i for i, child in enumerate(children[first_table_idx + 1 :], start=first_table_idx + 1)
            if child.tag.endswith("}p") and _element_text(child).startswith("Ref No")
        ),
        first_table_idx,
    ) + 1

    for block in extra_blocks:
        block_nodes = [_build_page_break_paragraph()] + [deepcopy(node) for node in block]
        for node in block_nodes:
            body.insert(insert_idx, node)
            insert_idx += 1

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.getvalue()

def convert_docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """
    Convierte un DOCX a PDF usando LibreOffice en modo headless.
    Preserva al máximo el layout porque renderiza el mismo archivo Word,
    sin transcribir ni reconstruir contenido.
    """
    soffice_bin = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice_bin:
        raise RuntimeError(
            "No encontré LibreOffice (soffice) en el servidor. "
            "Instálalo para habilitar la exportación exacta a PDF."
        )

    with tempfile.TemporaryDirectory() as tmp_dir:
        input_path = Path(tmp_dir) / "documento.docx"
        output_dir = Path(tmp_dir) / "out"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / "documento.pdf"
        input_path.write_bytes(docx_bytes)

        cmd = [
            soffice_bin,
            "--headless",
            "--convert-to",
            "pdf:writer_pdf_Export",
            str(input_path),
            "--outdir",
            str(output_dir),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(
                "La conversión a PDF falló en LibreOffice. "
                f"Detalle: {(result.stderr or result.stdout).strip()}"
            )
        if not output_path.exists():
            raise RuntimeError("LibreOffice no generó el PDF esperado.")
        return output_path.read_bytes()
#############################################################################################################################################################################
#############################################################################################################################################################################

def construir_plan_liquidacion(cronograma_df: pd.DataFrame, comision_mensual: float) -> pd.DataFrame:
    if cronograma_df.empty:
        return pd.DataFrame(columns=[
            "plan_key",
            "Fecha Límite de Pago",
            "Pago a Banco",
            "Comisión de Éxito",
            "Comisión Mensual",
            "Apartado Requerido",
        ])

    df = cronograma_df[cronograma_df["Cantidad"] > 0.005].copy()
    df["Fecha"] = pd.to_datetime(df["Fecha"])
    df["periodo"] = df["Fecha"].dt.to_period("M").astype(str)

    filas = []
    for periodo, group in df.groupby("periodo", sort=True):
        fecha_limite = group["Fecha"].min().date()
        pago_banco_mes = float(group.loc[group["Concepto"].str.contains("Entidad Financiera", na=False), "Cantidad"].sum())
        comision_exito_mes = float(group.loc[group["Concepto"].str.contains("Comisión Resuelve", na=False), "Cantidad"].sum())
        comision_mensual_mes = float(comision_mensual or 0.0)
        filas.append({
            "plan_key": periodo,
            "Fecha Límite de Pago": fecha_limite,
            "Pago a Banco": pago_banco_mes,
            "Comisión de Éxito": comision_exito_mes,
            "Comisión Mensual": comision_mensual_mes,
            "Apartado Requerido": pago_banco_mes + comision_exito_mes + comision_mensual_mes,
        })

    return pd.DataFrame(filas)


def construir_cronograma_pagos(
    fecha_inicial: date,
    plazo: int,
    n_pab: int,
    pago_banco_total: float,
    primer_pago_banco: float,
    comision_total: float,
    comision_inicial: float,
    dia_pago_banco: int | None = None,
    dia_pago_comision: int | None = None,
):
    plazo = max(int(plazo), 0)
    n_pab = max(int(n_pab), 1)
    primer_pago_banco = min(max(float(primer_pago_banco or 0.0), 0.0), max(float(pago_banco_total or 0.0), 0.0))
    pago_banco_total = max(float(pago_banco_total or 0.0), 0.0)
    comision_total = max(float(comision_total or 0.0), 0.0)
    comision_inicial = min(max(float(comision_inicial or 0.0), 0.0), comision_total)

    banco_restante = max(0.0, pago_banco_total - primer_pago_banco)
    meses_banco_restantes = max(0, n_pab - 1)
    meses_comision_restantes = max(0, plazo - meses_banco_restantes)
    comision_restante = max(0.0, comision_total - comision_inicial)

    pagos_banco = [primer_pago_banco]
    if meses_banco_restantes > 0:
        pagos_banco += _sum_rounded_parts([banco_restante / meses_banco_restantes] * meses_banco_restantes)

    pagos_comision = [comision_inicial]
    if meses_comision_restantes > 0:
        pagos_comision += _sum_rounded_parts([comision_restante / meses_comision_restantes] * meses_comision_restantes)

    filas = []
    if pagos_banco[0] > 0:
        filas.append({"Fecha": fecha_inicial, "Cantidad": pagos_banco[0], "Concepto": "Pago 1 a Entidad Financiera", "tipo": "banco", "orden": 0, "months_ahead": 0, "row_key": "banco_0"})
    if pagos_comision[0] > 0:
        filas.append({"Fecha": fecha_inicial, "Cantidad": pagos_comision[0], "Concepto": "Comisión Resuelve", "tipo": "comision", "orden": 1, "months_ahead": 0, "row_key": "comision_0"})

    for idx, valor in enumerate(pagos_banco[1:], start=1):
        if valor <= 0:
            continue
        filas.append({"Fecha": _day_of_month(fecha_inicial, idx, dia_pago_banco), "Cantidad": valor, "Concepto": f"Pago {idx + 1} a Entidad Financiera", "tipo": "banco", "orden": len(filas), "months_ahead": idx, "row_key": f"banco_{idx}"})

    for idx, valor in enumerate(pagos_comision[1:], start=1):
        if valor <= 0:
            continue
        offset = meses_banco_restantes + idx
        filas.append({"Fecha": _day_of_month(fecha_inicial, offset, dia_pago_comision), "Cantidad": valor, "Concepto": "Comisión Resuelve", "tipo": "comision", "orden": len(filas), "months_ahead": offset, "row_key": f"comision_{idx}"})

    cronograma = pd.DataFrame(filas)
    if cronograma.empty:
        cronograma = pd.DataFrame(columns=["Fecha", "Cantidad", "Concepto", "tipo", "orden", "months_ahead", "row_key"])
    return cronograma, {"meses_banco_restantes": meses_banco_restantes, "meses_comision_restantes": meses_comision_restantes}


def aplicar_overrides_cronograma(
    cronograma_df: pd.DataFrame,
    overrides_map: dict,
    totales_por_tipo: dict,
    fecha_inicial: date,
    dia_pago_banco: int | None,
    dia_pago_comision: int | None,
    primer_pago_banco_input: float,
    comision_inicial_input: float,
    lock_initial_rows: bool = True,
):
    if cronograma_df.empty:
        return cronograma_df.copy(), []

    df = cronograma_df.copy()
    df["cantidad_editada"] = False
    df["fecha_editada"] = False
    advertencias = []

    for row_key, cambios in (overrides_map or {}).items():
        matches = df.index[df["row_key"] == row_key].tolist()
        if not matches:
            continue
        idx = matches[0]
        if lock_initial_rows and int(df.at[idx, "months_ahead"]) == 0:
            continue
        if "Fecha" in cambios and cambios["Fecha"]:
            try:
                df.at[idx, "Fecha"] = pd.to_datetime(cambios["Fecha"]).date()
                df.at[idx, "fecha_editada"] = True
            except Exception:
                advertencias.append(f"No pude interpretar la fecha editada de {row_key}.")
        if "Cantidad" in cambios:
            try:
                df.at[idx, "Cantidad"] = max(_parse_amount_input(cambios["Cantidad"]), 0.0)
                df.at[idx, "cantidad_editada"] = True
            except Exception:
                advertencias.append(f"No pude interpretar el monto editado de {row_key}.")

    if lock_initial_rows:
        for tipo, input_value in [("banco", primer_pago_banco_input), ("comision", comision_inicial_input)]:
            mask = (df["tipo"] == tipo) & (df["months_ahead"] == 0)
            if mask.any():
                idx = df.index[mask][0]
                df.at[idx, "Cantidad"] = max(float(input_value or 0.0), 0.0)
                df.at[idx, "cantidad_editada"] = True

    partes = []
    for tipo, group in df.groupby("tipo", sort=False):
        total_objetivo = totales_por_tipo.get(tipo, float(group["Cantidad"].sum()))
        partes.append(_rebalance_group_amounts(group, total_objetivo))
    df = pd.concat(partes).sort_values("orden").reset_index(drop=True)

    banco_mask = (df["tipo"] == "banco") & (df["months_ahead"] > 0)
    banco_ocupados = set()
    for idx in df.index[banco_mask]:
        if bool(df.at[idx, "fecha_editada"]):
            banco_ocupados.add(max(1, _month_offset(fecha_inicial, pd.to_datetime(df.at[idx, "Fecha"]).date())))
        else:
            offset = int(df.at[idx, "months_ahead"])
            banco_ocupados.add(offset)
            df.at[idx, "Fecha"] = _day_of_month(fecha_inicial, offset, dia_pago_banco)

    comision_mask = (df["tipo"] == "comision") & (df["months_ahead"] > 0)
    offsets_fijos = set()
    for idx in df.index[comision_mask]:
        if bool(df.at[idx, "fecha_editada"]):
            offsets_fijos.add(max(1, _month_offset(fecha_inicial, pd.to_datetime(df.at[idx, "Fecha"]).date())))

    siguiente_offset = 1
    for idx in df.index[comision_mask]:
        if bool(df.at[idx, "fecha_editada"]):
            continue
        while siguiente_offset in offsets_fijos:
            siguiente_offset += 1
        while siguiente_offset in banco_ocupados:
            siguiente_offset += 1
        df.at[idx, "months_ahead"] = siguiente_offset
        df.at[idx, "Fecha"] = _day_of_month(fecha_inicial, siguiente_offset, dia_pago_comision)
        offsets_fijos.add(siguiente_offset)
        siguiente_offset += 1

    df = df.sort_values(by=["Fecha", "orden"]).reset_index(drop=True)
    return df, advertencias
#####################################################################################################################################################################################    

def _model_version() -> str:
    return _file_version(MODEL_PATH)

def _meta_version() -> str:
    return _file_version(META_PATH)

def _data_version() -> str:
    # Si existe parquet, usamos ese; si no, miramos CSV.
    if DATA_PARQUET.exists():
        return _file_version(DATA_PARQUET)
    if DATA_CSV.exists():
        return _file_version(DATA_CSV)
    return "missing"

def _parse_relaxed_service_account_string(raw_value: str) -> dict:
    """
    Intenta reconstruir un JSON de service account aunque llegue como texto
    con saltos reales dentro de `private_key`, wrappers extra o texto copiado
    desde distintos gestores de secretos.
    """
    required_fields = ["private_key", "client_email"]
    known_fields = [
        "type",
        "project_id",
        "private_key_id",
        "private_key",
        "client_email",
        "client_id",
        "auth_uri",
        "token_uri",
        "auth_provider_x509_cert_url",
        "client_x509_cert_url",
        "universe_domain",
    ]

    text = raw_value.strip()

    if text.startswith("MI_JSON="):
        text = text.split("=", 1)[1].strip()

    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        text = text[start : end + 1]

    parsed = {}
    for i, field in enumerate(known_fields):
        field_match = re.search(rf'"{re.escape(field)}"\s*:\s*"', text)
        if not field_match:
            continue

        value_start = field_match.end()
        next_positions = []
        for next_field in known_fields[i + 1 :]:
            next_match = re.search(rf'",\s*"{re.escape(next_field)}"\s*:', text[value_start:])
            if next_match:
                next_positions.append(value_start + next_match.start())

        if next_positions:
            value_end = min(next_positions)
        else:
            tail_match = re.search(r'"\s*}$', text[value_start:])
            if not tail_match:
                continue
            value_end = value_start + tail_match.start()

        parsed[field] = text[value_start:value_end]

    if not all(field in parsed for field in required_fields):
        raise RuntimeError(
            "El secreto `MI_JSON` no se pudo interpretar. Revisa que tenga el JSON completo del service account."
        )

    return parsed
    
def _looks_like_service_account_mapping(value) -> bool:
    try:
        data = dict(value)
    except Exception:
        return False
    return "private_key" in data and "client_email" in data

def _extract_service_account_from_secrets_tree(value):
    """
    Busca credenciales en estructuras anidadas de st.secrets (dict/AttrDict).
    Soporta:
    - llave MI_JSON en cualquier nivel
    - objetos con private_key + client_email en cualquier nivel
    """
    try:
        if isinstance(value, str) and value.strip():
            return value
    except Exception:
        pass

    if _looks_like_service_account_mapping(value):
        return dict(value)

    try:
        items = dict(value).items()
    except Exception:
        return None

    for key, sub_value in items:
        if str(key).strip().upper() == "MI_JSON":
            return sub_value
        nested = _extract_service_account_from_secrets_tree(sub_value)
        if nested is not None:
            return nested
    return None



def _load_google_service_account_info() -> dict:
    """
    Carga el JSON del service account desde Streamlit Secrets o variable de entorno.
    Soporta MI_JSON como string JSON, tabla TOML, dict directo o variables separadas.
    """
    creds_source = None

    try:
        creds_source = _extract_service_account_from_secrets_tree(st.secrets)    
    except Exception:
        creds_source = None

    if creds_source is None:
        env_json = (
            os.environ.get("MI_JSON")
            or os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON")
            or os.environ.get("MI_JSON_SANDBOX")
        )
        if env_json:
            creds_source = env_json
        else:
            env_fields = {
                key: os.environ.get(key)
                for key in [
                    "type",
                    "project_id",
                    "private_key_id",
                    "private_key",
                    "client_email",
                    "client_id",
                    "auth_uri",
                    "token_uri",
                    "auth_provider_x509_cert_url",
                    "client_x509_cert_url",
                    "universe_domain",
                ]
                if os.environ.get(key) is not None
            }
            if _looks_like_service_account_mapping(env_fields):
                creds_source = env_fields

    if creds_source is None:
        raise RuntimeError(
            "No encontré credenciales de Google Sheets. "
            "Configura `MI_JSON` en Streamlit Secrets (de ESTE despliegue) "
            "o una variable de entorno (`MI_JSON` / `GOOGLE_SERVICE_ACCOUNT_JSON`)."
        )

    if isinstance(creds_source, str):
        creds_source = creds_source.strip()
        if not creds_source:
            raise RuntimeError("El secreto `MI_JSON` está vacío.")
        try:
            creds_info = json.loads(creds_source)
        except json.JSONDecodeError:
            try:
                creds_info = ast.literal_eval(creds_source)
            except Exception:
                creds_info = _parse_relaxed_service_account_string(creds_source)
    
    else:
        try:
            creds_info = dict(creds_source)
        except Exception as exc:
            raise RuntimeError("No pude interpretar el secreto `MI_JSON` como credenciales válidas.") from exc

    private_key = creds_info.get("private_key")
    if isinstance(private_key, str):
        normalized_key = private_key.strip().replace("\r\n", "\n").replace("\\n", "\n")
        if "-----BEGIN PRIVATE KEY-----" in normalized_key and not normalized_key.endswith("\n"):
            normalized_key += "\n"
        creds_info["private_key"] = normalized_key

    return creds_info


@st.cache_resource(show_spinner=False)
def get_google_sheet_worksheet(tab_name: str = GOOGLE_SHEET_TAB):
    """
    Devuelve la hoja de cálculo destino para histórico.
    Se cachea mientras no cambie el proceso.
    """
    creds_info = _load_google_service_account_info()
    credentials = Credentials.from_service_account_info(creds_info, scopes=GOOGLE_SHEETS_SCOPES)
    client = gspread.authorize(credentials)
    spreadsheet = client.open_by_key(GOOGLE_SHEET_ID)
    return spreadsheet.worksheet(tab_name)


def _append_row_to_google_sheet(row_data: dict):
    """
    Inserta una fila en Google Sheets y retorna (ok, destination, error_msg).
    """
    try:
        worksheet = get_google_sheet_worksheet()
        expected_headers = GOOGLE_SHEET_HEADERS
        current_headers = worksheet.row_values(1)
        if current_headers[: len(expected_headers)] != expected_headers:
            worksheet.update("A1:H1", [expected_headers])

        worksheet.append_row(
            [row_data.get(header, "") for header in expected_headers],
            value_input_option="USER_ENTERED",
        )
        return True, f"Google Sheets > {GOOGLE_SHEET_TAB}", None
    except Exception as e:
        return False, f"Google Sheets > {GOOGLE_SHEET_TAB}", str(e)

def _col_index_to_letter(col_idx: int) -> str:
    letters = ""
    while col_idx > 0:
        col_idx, rem = divmod(col_idx - 1, 26)
        letters = chr(65 + rem) + letters
    return letters or "A"


def _append_row_to_respuestas_estr(row_data: dict):
    """
    Inserta una fila en la hoja "Respuestas Estr" mapeando por encabezado visible.
    Evita desalineaciones cuando la estructura de columnas cambia en la hoja.
    """
    try:
        worksheet = get_google_sheet_worksheet(GOOGLE_SHEET_TAB_RESPUESTAS)
        headers = worksheet.row_values(1)
        if not headers:
            headers = GOOGLE_RESPUESTAS_COLS.copy()

        normalized = [""] * len(headers)
        payload = dict(row_data or {})

        for idx, header in enumerate(headers):
            h = _norm(header)
            if "marca temporal" in h:
                normalized[idx] = payload.get("timestamp", "")
            elif "direccion de correo" in h or h == "correo":
                normalized[idx] = payload.get("correo_electronico", "")
            elif h == "referencia":
                normalized[idx] = payload.get("referencia", "")
            elif "id deuda" in h or "id de la deuda" in h or "id de deuda" in h:
                normalized[idx] = payload.get("ids", "")
            elif "banco" in h:
                normalized[idx] = payload.get("bancos", "")
            elif "carta" in h and "pagare" in h:
                normalized[idx] = payload.get("carta_pagare_link", "")
            elif "pantallazo" in h and "aceptacion" in h:
                normalized[idx] = payload.get("pantallazo_aceptacion_link", "")
            elif "condonacion" in h and "mensualidades" in h:
                normalized[idx] = payload.get("condonacion_mensualidades", "")
            elif "pantallazo" in h and "correo" in h and "condonacion" in h:
                normalized[idx] = payload.get("pantallazo_correo_condonacion_link", "")
            elif "total de comision" in h:
                normalized[idx] = payload.get("comision_exito_total", "")
            elif "primera comision" in h or "pago de la primera comision" in h:
                normalized[idx] = payload.get("ce_inicial", "")
            elif "aprobacion estructurados" in h:
                normalized[idx] = payload.get("es_aprobado_bool", "")
            elif h == "estado" or "comentario" in h:
                normalized[idx] = payload.get("estado_aprobacion", "")
            elif "calculadora" in h:
                normalized[idx] = payload.get("prediccion", "")

        # Regla de negocio: la última columna debe registrar la predicción
        # usada para enviar a aprobación (control anti-manipulación).
        if normalized:
            normalized[-1] = payload.get("prediccion", "")

        # Evita que el registro se vaya al final de filas "ocupadas" por fórmulas
        # (por ejemplo, columnas con FALSE/checkbox). Lo insertamos justo después
        # de la última fila donde A, B y C están diligenciadas.
        col_a = worksheet.col_values(1)
        col_b = worksheet.col_values(2)
        col_c = worksheet.col_values(3)
        max_len = max(len(col_a), len(col_b), len(col_c), 1)

        last_data_row = 1  # encabezados
        for row_idx in range(2, max_len + 1):
            a_val = col_a[row_idx - 1] if row_idx <= len(col_a) else ""
            b_val = col_b[row_idx - 1] if row_idx <= len(col_b) else ""
            c_val = col_c[row_idx - 1] if row_idx <= len(col_c) else ""
            if str(a_val).strip() and str(b_val).strip() and str(c_val).strip():
                last_data_row = row_idx

        target_row = last_data_row + 1
        end_col_letter = _col_index_to_letter(len(headers))
        worksheet.update(
            f"A{target_row}:{end_col_letter}{target_row}",
            [normalized],
            value_input_option="USER_ENTERED",
        )
        return True, f"Google Sheets > {GOOGLE_SHEET_TAB_RESPUESTAS}", None
    except Exception as e:
        return False, f"Google Sheets > {GOOGLE_SHEET_TAB_RESPUESTAS}", str(e)

def _parse_sheet_timestamp(value: str) -> datetime | None:
    text = str(value or "").strip()
    if not text:
        return None
    for fmt in ("%d/%m/%Y %H:%M:%S", "%d/%m/%Y %H:%M", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue
    return None


def _ids_to_set(ids_value) -> set[str]:
    if isinstance(ids_value, (list, tuple, set)):
        raw = [str(x) for x in ids_value]
    else:
        raw = re.split(r"[^0-9A-Za-z]+", str(ids_value or ""))
    return {item.strip() for item in raw if str(item).strip()}


def _get_respuestas_duplicados_mes(referencia, ids) -> dict:
    """
    Busca duplicados del mes en "Respuestas Estr":
      - exact_duplicate: coincide referencia + set de IDs
      - reference_duplicate: coincide referencia, pero IDs distintos
    """
    try:
        worksheet = get_google_sheet_worksheet(GOOGLE_SHEET_TAB_RESPUESTAS)
        all_values = worksheet.get_all_values()
        if not all_values:
            return {"ok": True, "mode": "none", "exact_rows": [], "error": None}

        headers = all_values[0]
        rows = all_values[1:]
        header_idx = { _norm(h): i for i, h in enumerate(headers) }

        timestamp_idx = next((i for h, i in header_idx.items() if "marca temporal" in h), None)
        ref_idx = next((i for h, i in header_idx.items() if h == "referencia"), None)
        ids_idx = next((i for h, i in header_idx.items() if "id deuda" in h or "id de la deuda" in h or "id de deuda" in h), None)
        aprob_idx = next((i for h, i in header_idx.items() if "aprobacion estructurados" in h), None)
        comentario_idx = next((i for h, i in header_idx.items() if h == "estado" or "comentario" in h), None)

        if timestamp_idx is None or ref_idx is None or ids_idx is None:
            return {"ok": True, "mode": "none", "exact_rows": [], "error": None}

        now = datetime.now()
        ref_norm = _norm(str(referencia or ""))
        ids_target = _ids_to_set(ids)
        ref_found = False
        exact_rows = []

        for offset, row in enumerate(rows, start=2):
            ts_raw = row[timestamp_idx] if timestamp_idx < len(row) else ""
            ref_raw = row[ref_idx] if ref_idx < len(row) else ""
            ids_raw = row[ids_idx] if ids_idx < len(row) else ""
            row_dt = _parse_sheet_timestamp(ts_raw)
            if row_dt is None:
                continue
            if row_dt.year != now.year or row_dt.month != now.month:
                continue
            if _norm(ref_raw) != ref_norm:
                continue
            ref_found = True
            if _ids_to_set(ids_raw) == ids_target:
                exact_rows.append(
                    {
                        "row_idx": offset,
                        "aprob_col_idx": aprob_idx + 1 if aprob_idx is not None else None,
                        "comentario_col_idx": comentario_idx + 1 if comentario_idx is not None else None,
                        "comentario_actual": row[comentario_idx] if comentario_idx is not None and comentario_idx < len(row) else "",
                    }
                )

        mode = "none"
        if exact_rows:
            mode = "exact_duplicate"
        elif ref_found:
            mode = "reference_duplicate"

        return {"ok": True, "mode": mode, "exact_rows": exact_rows, "error": None}
    except Exception as e:
        return {"ok": False, "mode": "none", "exact_rows": [], "error": str(e)}


def _marcar_anteriores_como_duplicado(exact_rows: list[dict]):
    if not exact_rows:
        return
    worksheet = get_google_sheet_worksheet(GOOGLE_SHEET_TAB_RESPUESTAS)
    for row_info in exact_rows:
        row_idx = row_info.get("row_idx")
        aprob_col_idx = row_info.get("aprob_col_idx")
        comentario_col_idx = row_info.get("comentario_col_idx")
        comentario_actual = _norm(row_info.get("comentario_actual", ""))
        if row_idx and aprob_col_idx:
            target_cell = f"{_col_index_to_letter(aprob_col_idx)}{row_idx}"
            worksheet.update(
                target_cell,
                [["FALSE"]],
                value_input_option="USER_ENTERED",
            )
        if row_idx and comentario_col_idx and comentario_actual == "aprobado":
            target_cell = f"{_col_index_to_letter(comentario_col_idx)}{row_idx}"
            worksheet.update(
                target_cell,
                [["Duplicado"]],
                value_input_option="USER_ENTERED",
            )

def diagnosticar_google_sheets():
    """
    Valida que el secreto, el spreadsheet y la pestaña destino estén accesibles.
    No escribe datos; solo devuelve el estado para mostrarlo en la UI.
    """
    try:
        creds_info = _load_google_service_account_info()
        client_email = creds_info.get("client_email", "desconocido")

        credentials = Credentials.from_service_account_info(creds_info, scopes=GOOGLE_SHEETS_SCOPES)
        client = gspread.authorize(credentials)
        spreadsheet = client.open_by_key(GOOGLE_SHEET_ID)
        worksheet = spreadsheet.worksheet(GOOGLE_SHEET_TAB)

        return {
            "ok": True,
            "client_email": client_email,
            "spreadsheet_title": spreadsheet.title,
            "worksheet_title": worksheet.title,
            "error": None,
        }
    except Exception as e:
        client_email = "desconocido"
        try:
            creds_info = _load_google_service_account_info()
            client_email = creds_info.get("client_email", "desconocido")
        except Exception:
            pass

        return {
            "ok": False,
            "client_email": client_email,
            "spreadsheet_title": None,
            "worksheet_title": None,
            "error": str(e),
        }


def google_sheets_status():
    """
    Alias simple para evitar errores por diferencias de nombre al invocar el diagnóstico.
    """
    return diagnosticar_google_sheets()
        
# =================== Loaders cacheados ===================
@st.cache_resource(show_spinner=False)
def load_model(_version: str):
    """
    Carga el modelo MLP. _version es la firma del archivo y se usa
    solo para invalidar el cache cuando cambie el .joblib.
    """
    if not MODEL_PATH.exists():
        st.error("No encontré el archivo del modelo `mlp_recaudo_pipeline.joblib` en la raíz del repo.")
        st.stop()
    return load(MODEL_PATH)

@st.cache_data(show_spinner=False)
def load_meta(_version: str):
    """
    Carga el JSON de metadata. Se refresca cuando cambie el archivo.
    """
    meta = {}
    if META_PATH.exists():
        try:
            meta = json.loads(META_PATH.read_text(encoding="utf-8"))
        except Exception:
            meta = {}
    return meta

@st.cache_data(show_spinner=False)
def load_repo_base(_version: str) -> pd.DataFrame | None:
    """
    Carga la base que deja el workflow.
    Si existen Parquet y CSV, usa el CSV cuando trae más columnas que el Parquet
    (por ejemplo, campos enriquecidos para poblar el documento Word).
    _version se usa solo para invalidar cache cuando cambian los archivos.
    Devuelve None si no existe.
    """
    try:
        df_parquet = pd.read_parquet(DATA_PARQUET) if DATA_PARQUET.exists() else None
        df_csv = pd.read_csv(DATA_CSV) if DATA_CSV.exists() else None

        if df_parquet is not None and df_csv is not None:
            return df_csv if len(df_csv.columns) > len(df_parquet.columns) else df_parquet
        if df_parquet is not None:
            return df_parquet
        if df_csv is not None:
            return df_csv
        return None
    except Exception:
        # Si algo falla leyendo, permitimos fallback a subida manual
        return None

@st.cache_data(show_spinner=False)
def load_clientes_lookup() -> pd.DataFrame | None:
    if not CLIENTES_LOOKUP_PATH.exists():
        return None

    def _clean_df(df: pd.DataFrame) -> pd.DataFrame:
        df = df.copy()
        df.columns = [str(col).replace("﻿", "").strip() for col in df.columns]
        for col in df.columns:
            df[col] = df[col].map(lambda value: str(value).replace("﻿", "").strip() if value is not None else "")
        return df

    def _read_with_sep(sep, engine=None):
        kwargs = {"encoding": "latin-1", "dtype": str, "keep_default_na": False}
        if sep is not None:
            kwargs["sep"] = sep
        if engine is not None:
            kwargs["engine"] = engine
        return pd.read_csv(CLIENTES_LOOKUP_PATH, **kwargs)

    try:
        sample = CLIENTES_LOOKUP_PATH.read_text(encoding="latin-1", errors="ignore")[:4096]
        dialect = csv.Sniffer().sniff(sample, delimiters=",;")
        df = _read_with_sep(dialect.delimiter)
    except Exception:
        try:
            df = _read_with_sep(None, engine="python")
        except Exception:
            return None

    if len(df.columns) == 1:
        only_col = str(df.columns[0])
        fallback_sep = ";" if ";" in only_col else ","
        try:
            df = _read_with_sep(fallback_sep)
        except Exception:
            return _clean_df(df)

    return _clean_df(df)


def _normalize_lookup_key(value) -> str:
    text_value = str(value or "").strip()
    if text_value.endswith(".0"):
        text_value = text_value[:-2]
    digits_only = re.sub(r"\D", "", text_value)
    if digits_only and len(digits_only) >= max(6, len(text_value.replace(" ", "")) - 2):
        return digits_only
    return text_value


def _format_city_department(ciudad, departamento) -> str:
    ciudad_text = _smart_title_case(ciudad)
    departamento_text = _smart_title_case(departamento)
    normalized_city = _norm(ciudad_text)
    if normalized_city in {"bogota d c", "bogota dc", "bogota"}:
        return "Bogotá D.C."
    if ciudad_text and departamento_text:
        return f"{ciudad_text}, {departamento_text}"
    return ciudad_text or departamento_text


def _lookup_cliente_info(referencia, cedula_cliente) -> dict[str, str]:
    clientes_df = load_clientes_lookup()
    if clientes_df is None or clientes_df.empty:
        return {}

    col_ref = _find_col(clientes_df, ["Referencia"])
    col_doc = _find_col(clientes_df, ["Documento"])
    col_cel = _find_col(clientes_df, ["Celular"])
    col_ciu = _find_col(clientes_df, ["Ciudad"])
    col_dep = _find_col(clientes_df, ["Departamento"])
    col_dir = _find_col(clientes_df, ["Direccion", "Dirección"])

    match = pd.DataFrame()
    ref_text = _normalize_lookup_key(referencia)
    cedula_text = _normalize_lookup_key(cedula_cliente)

    if col_ref and ref_text:
        match = clientes_df[clientes_df[col_ref].map(_normalize_lookup_key) == ref_text]
    if match.empty and col_doc and cedula_text:
        match = clientes_df[clientes_df[col_doc].map(_normalize_lookup_key) == cedula_text]
    if match.empty:
        return {}

    row = match.iloc[0]
    return {
        "telefono_cliente": str(row[col_cel]).strip() if col_cel and pd.notna(row[col_cel]) else "",
        "ciudad_cliente": _format_city_department(row[col_ciu] if col_ciu else "", row[col_dep] if col_dep else ""),
        "direccion_cliente": str(row[col_dir]).strip() if col_dir and pd.notna(row[col_dir]) else "",
    }
# =================== LOG LOCAL ===================
def _get_writable_log_path() -> Path:
    """
    Devuelve la ruta donde se guardará el histórico.
    Si `data/` no es escribible (por despliegue), usa un fallback temporal.
    """
    candidates = [
        LOG_PATH,
        Path(gettempdir()) / "recaudo-app" / "logs_calculadora.csv",
    ]

    for path in candidates:
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            with path.parent.joinpath(".write_test").open("w", encoding="utf-8") as f:
                f.write("ok")
            path.parent.joinpath(".write_test").unlink(missing_ok=True)
            return path
        except Exception:
            continue

    return LOG_PATH


def guardar_log_calculo(
    referencia,
    ids,
    features,
    prediccion,
):
    """
    Guarda una fila del histórico en Google Sheets y, como respaldo,
    también en CSV local. Retorna un diccionario con el resultado.
    """
    
    fila = {
        "fecha": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "referencia": str(referencia),
        "ids_deuda": ",".join(map(str, ids)),
        "plazo": features.get("PRI-ULT"),
        "ratio_pp": features.get("Ratio_PP"),
        "cuota_apartado": features.get("C/A"),
        "amount_total": features.get("AMOUNT_TOTAL"),
        "prediccion": prediccion,
    }
    
    sheet_ok, sheet_dest, sheet_err = _append_row_to_google_sheet(fila)

    log_path = _get_writable_log_path()
    local_ok = False
    local_err = None
    try:
        df_log = pd.DataFrame([fila])
        file_exists = log_path.exists() and log_path.stat().st_size > 0
        df_log.to_csv(
            log_path,
            mode="a" if file_exists else "w",
            header=not file_exists,
            index=False,
            encoding="utf-8-sig",
        )
        local_ok = True
    except Exception as e:
        local_err = str(e)

    return {
        "sheet_ok": sheet_ok,
        "sheet_destination": sheet_dest,
        "sheet_error": sheet_err,
        "local_ok": local_ok,
        "local_path": log_path,
        "local_error": local_err,
    }

def enviar_aprobacion_estructurados(
    *,
    referencia,
    ids,
    bancos,
    correo_electronico,
    condonacion_mensualidades,
    comision_exito_total,
    ce_inicial,
    prediccion,
    tipo_liquidacion="",
    carta_pagare_link="Pendiente",
    pantallazo_aceptacion_link="Pendiente",
    pantallazo_correo_condonacion_link="Pendiente",
    duplicate_mode="none",
    exact_rows_previas=None,
    
):
    tipo_liquidacion_norm = _norm(tipo_liquidacion)
    umbral_aprobacion = 0.8 if "tradicional" in tipo_liquidacion_norm else 0.74
    es_aprobado = float(prediccion or 0.0) >= float(umbral_aprobacion)
    condonacion_value = "Sí" if str(condonacion_mensualidades).strip().lower() == "si" else "No"
    correo_condonacion_link_value = (
        str(pantallazo_correo_condonacion_link or "").strip()
        if condonacion_value == "Sí"
        else ""
    )

    if duplicate_mode == "reference_duplicate":
        es_aprobado_bool = ""
        estado_aprobacion = ""
    else:
        es_aprobado_bool = "TRUE" if es_aprobado else "FALSE"
        estado_aprobacion = "Aprobado" if es_aprobado else "Rechazado"
    

    respuestas_payload = {
        "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "correo_electronico": str(correo_electronico or "").strip(),
        "referencia": str(referencia),
        "ids": "-".join(map(str, ids)),
        "bancos": str(bancos),
        "carta_pagare_link": str(carta_pagare_link or "").strip(),
        "pantallazo_aceptacion_link": str(pantallazo_aceptacion_link or "").strip(),
        "condonacion_mensualidades": condonacion_value,
        "pantallazo_correo_condonacion_link": correo_condonacion_link_value,
        "comision_exito_total": float(comision_exito_total or 0.0),
        "ce_inicial": float(ce_inicial or 0.0),
        "es_aprobado_bool": es_aprobado_bool,
        "estado_aprobacion": estado_aprobacion,
        "prediccion": float(prediccion or 0.0),
    }
    if duplicate_mode == "exact_duplicate":
        _marcar_anteriores_como_duplicado(exact_rows_previas or [])
    estr_ok, estr_dest, estr_err = _append_row_to_respuestas_estr(respuestas_payload)
    return {
        "estr_ok": estr_ok,
        "estr_destination": estr_dest,
        "estr_error": estr_err,
        "es_aprobado": es_aprobado,
        "umbral_aprobacion": umbral_aprobacion,
        "duplicate_mode": duplicate_mode,
    }
        

# ------------------ utilidades ------------------
def _norm(s: str) -> str:
    # ✅ MEJORADO: soporta guiones, underscores, espacios raros, etc.
    s = str(s).strip().lower()
    rep = str.maketrans("áéíóúü", "aeiouu")
    s = s.translate(rep)
    s = s.replace("\xa0", " ")
    s = s.replace("_", " ").replace("-", " ")
    s = re.sub(r"[^a-z0-9 ]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _load_google_oauth_client_config() -> dict:
    """
    Carga client config OAuth para usuario final.
    Busca en secrets:
      - GOOGLE_OAUTH_CLIENT_JSON (JSON completo de OAuth Client)
      - o GOOGLE_OAUTH_CLIENT_ID + GOOGLE_OAUTH_CLIENT_SECRET
    """
    client_id = st.secrets.get("GOOGLE_OAUTH_CLIENT_ID")
    client_secret = st.secrets.get("GOOGLE_OAUTH_CLIENT_SECRET")
    if client_id and client_secret:
        return {
            "installed": {
                "client_id": str(client_id),
                "client_secret": str(client_secret),
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "redirect_uris": ["http://localhost"],
            }
        }

    raw_json = st.secrets.get("GOOGLE_OAUTH_CLIENT_JSON")
    if raw_json:
        if isinstance(raw_json, dict):
            cfg = raw_json
        else:
            cfg = json.loads(str(raw_json))
        if "installed" in cfg or "web" in cfg:
            return cfg
        if all(k in cfg for k in ("client_id", "client_secret", "auth_uri", "token_uri")):
            return {"installed": cfg}
        if str(cfg.get("type", "")).strip().lower() == "service_account":
            raise ValueError(
                "GOOGLE_OAUTH_CLIENT_JSON parece ser de Service Account. "
                "Para OAuth de usuario debes usar credenciales de tipo Web/Installed App."
            )
        raise ValueError("GOOGLE_OAUTH_CLIENT_JSON no tiene formato OAuth válido.")
    raise RuntimeError(
        "Faltan secretos OAuth de Google. Configura GOOGLE_OAUTH_CLIENT_JSON o "
        "GOOGLE_OAUTH_CLIENT_ID + GOOGLE_OAUTH_CLIENT_SECRET."
    )


def _oauth_drive_configurado() -> bool:
    try:
        _load_google_oauth_client_config()
        return True
    except Exception:
        return False


def _get_oauth_redirect_uri() -> str:
    """
    URI de retorno OAuth.
    - Si existe GOOGLE_OAUTH_REDIRECT_URI en secrets, se usa ese valor.
    - Si no, se mantiene localhost para flujo manual.
    """
    custom_redirect = str(st.secrets.get("GOOGLE_OAUTH_REDIRECT_URI", "")).strip()
    return custom_redirect or "http://localhost"


def _get_drive_flow_and_auth_url():
    client_config = _load_google_oauth_client_config()
    flow = Flow.from_client_config(client_config, scopes=GOOGLE_DRIVE_UPLOAD_SCOPES)
    flow.redirect_uri = _get_oauth_redirect_uri()
    auth_url, state = flow.authorization_url(
        access_type="offline",
        include_granted_scopes="true",
        prompt="consent",
    )
    return flow, auth_url, state, client_config


def _extract_oauth_code(redirect_text: str) -> str:
    text = str(redirect_text or "").strip()
    if not text:
        return ""
    if text.startswith("http://") or text.startswith("https://"):
        parsed = urlparse(text)
        q = parse_qs(parsed.query)
        return str((q.get("code") or [""])[0]).strip()
    return text


def _build_drive_service_from_session():
    token_data = st.session_state.get("drive_user_token")
    if not token_data:
        return None
    creds = UserCredentials.from_authorized_user_info(token_data, GOOGLE_DRIVE_UPLOAD_SCOPES)
    if not creds.valid:
        return None
    return build("drive", "v3", credentials=creds)


def _upload_file_to_drive(
    service,
    uploaded_file,
    folder_id: str,
    *,
    allowed_extensions: tuple[str, ...],
    invalid_message: str,
):
    if uploaded_file is None:
        return None
    file_name = str(uploaded_file.name or "").strip()
    if not file_name:
        raise ValueError(invalid_message)
    lower_file_name = file_name.lower()
    if not any(lower_file_name.endswith(ext) for ext in allowed_extensions):
        raise ValueError(invalid_message)

    content = uploaded_file.getvalue()
    metadata = {"name": file_name, "parents": [folder_id]}
    mime_type = str(getattr(uploaded_file, "type", "") or "").strip() or "application/octet-stream"
    media = MediaIoBaseUpload(BytesIO(content), mimetype=mime_type, resumable=True)
    result = service.files().create(
        body=metadata,
        media_body=media,
        fields="id,name,webViewLink,parents",
        supportsAllDrives=True,
    ).execute()
    return result
def _normalize_reference_token(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9]", "", str(value or "")).upper().strip()


def _extract_reference_from_text(text: str) -> str:
    if not text:
        return ""
    patterns = [
        r"ref(?:erencia)?\s*(?:no|nro|n\.°|n°|num(?:ero)?)?[:.\-\s]*([A-Za-z0-9\-]{6,})",
        r"referencia[:.\-\s]*([A-Za-z0-9\-]{6,})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return _normalize_reference_token(match.group(1))
    return ""


def _extract_pdf_first_last_text(pdf_bytes: bytes) -> tuple[str, str]:
    first_page_text = ""
    last_page_text = ""

    pypdf_spec = importlib.util.find_spec("pypdf")
    if pypdf_spec:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(pdf_bytes))
        if len(reader.pages) > 0:
            first_page_text = str(reader.pages[0].extract_text() or "")
            last_page_text = str(reader.pages[-1].extract_text() or "")
        return first_page_text, last_page_text

    pypdf2_spec = importlib.util.find_spec("PyPDF2")
    if pypdf2_spec:
        from PyPDF2 import PdfReader

        reader = PdfReader(BytesIO(pdf_bytes))
        if len(reader.pages) > 0:
            first_page_text = str(reader.pages[0].extract_text() or "")
            last_page_text = str(reader.pages[-1].extract_text() or "")
        return first_page_text, last_page_text

    raise RuntimeError("No hay librería instalada para leer PDFs (pypdf/PyPDF2).")


def _has_qr_signal_in_last_page(last_page_text: str) -> bool:
    text_norm = str(last_page_text or "").lower()
    qr_signals = [
        "qr",
        "código qr",
        "codigo qr",
        "escanee",
        "autentic",
        "firma electrón",
        "firma electron",
    ]
    return any(signal in text_norm for signal in qr_signals)


def _extract_flow_rows_from_text(first_page_text: str) -> list[dict]:
    rows: list[dict] = []
    if not first_page_text:
        return rows

    lines = [re.sub(r"\s+", " ", str(line or "")).strip() for line in str(first_page_text).splitlines()]
    lines = [line for line in lines if line]
    row_regex = re.compile(r"^\s*(\d{1,3})\s+\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b")
    amount_with_cop_regex = re.compile(r"(\$?\s*\d{1,3}(?:[.,]\d{3})+(?:[.,]\d{2})?)\s*COP", flags=re.IGNORECASE)
    grouped_amount_regex = re.compile(r"(\$?\s*\d{1,3}(?:[.,]\d{3})+(?:[.,]\d{2})?)")

    for line in lines:
        row_match = row_regex.search(line)
        if not row_match:
            continue
        n_value = int(row_match.group(1))

        line_after_date = line[row_match.end():]
        amount_match = amount_with_cop_regex.search(line_after_date)
        amount_candidates = [amount_match.group(1)] if amount_match else grouped_amount_regex.findall(line_after_date)
        parsed_amount = None
        for candidate in amount_candidates:
            parsed = _parse_amount_input(candidate)
            if parsed > 0:
                parsed_amount = round(float(parsed), 2)
                break

        if parsed_amount is None:
            continue

        rows.append({"n": n_value, "amount": parsed_amount})
    return rows


def _build_expected_flow_rows(expected_flow_df: pd.DataFrame | None) -> list[dict]:
    if expected_flow_df is None or expected_flow_df.empty:
        return []
    required_cols = {"Fecha", "Cantidad", "Concepto"}
    if not required_cols.issubset(set(expected_flow_df.columns)):
        return []

    rows: list[dict] = []
    for idx, (_, row) in enumerate(expected_flow_df[["Cantidad"]].iterrows(), start=1):
        amount = round(float(pd.to_numeric(row["Cantidad"], errors="coerce") or 0.0), 2)
        rows.append({"n": int(idx), "amount": amount})
    return rows


def _validate_flow_matches_pdf(first_page_text: str, expected_flow_df: pd.DataFrame | None) -> tuple[bool, str]:
    expected_rows = _build_expected_flow_rows(expected_flow_df)
    if not expected_rows:
        return True, ""

    pdf_rows = _extract_flow_rows_from_text(first_page_text)
    if not pdf_rows:
        return False, "La tabla de flujo del PDF no se pudo leer correctamente."

    expected_last_n = max(int(r["n"]) for r in expected_rows) if expected_rows else 0
    pdf_last_n = max(int(r["n"]) for r in pdf_rows) if pdf_rows else 0
    if expected_last_n != pdf_last_n:
        return False, "La columna N de la tabla no coincide entre calculadora y PDF."

    expected_by_n = {int(r["n"]): float(r["amount"]) for r in expected_rows}
    pdf_by_n: dict[int, float] = {}
    for row in sorted(pdf_rows, key=lambda x: int(x["n"])):
        n_val = int(row["n"])
        # Tomamos la primera ocurrencia por N (tabla principal) y evitamos
        # que otras tablas de la página sobrescriban el valor.
        if n_val not in pdf_by_n:
            pdf_by_n[n_val] = float(row["amount"])

    if len(expected_by_n) != len(pdf_by_n):
        return False, "La columna Cantidad de la tabla no coincide entre calculadora y PDF."

    diferencias = []
    for n_val in sorted(expected_by_n.keys()):
        expected_amount = expected_by_n.get(n_val)
        pdf_amount = pdf_by_n.get(n_val)
        if pdf_amount is None:
            diferencias.append((n_val, expected_amount, None))
            continue
        if abs(float(expected_amount) - float(pdf_amount)) > 1.0:
            diferencias.append((n_val, expected_amount, pdf_amount))

    if diferencias:
        def _fmt_cop_plain(value: float | int) -> str:
            base = f"{float(value):,.2f}"
            base = base.replace(",", "X").replace(".", ",").replace("X", ".")
            return f"{base} COP"

        n_val, expected_amount, pdf_amount = diferencias[0]
        if pdf_amount is None:
            return False, f"La columna Cantidad no coincide: falta la fila N° {n_val} en la tabla del PDF."
        return (
            False,
            "La columna Cantidad no coincide en la fila "
            f"N° {n_val} (calculadora: {_fmt_cop_plain(expected_amount)} vs PDF: {_fmt_cop_plain(pdf_amount)})."
        )

    return True, ""


def _validate_carta_pagare_pdf(uploaded_file, expected_reference: str, expected_flow_df: pd.DataFrame | None = None) -> tuple[bool, str]:
    if uploaded_file is None:
        return False, "Debes adjuntar Carta/Pagaré (solo PDF)."
    try:
        pdf_bytes = uploaded_file.getvalue()
        first_page_text, last_page_text = _extract_pdf_first_last_text(pdf_bytes)
    except Exception as exc:
        return False, f"No pude leer el PDF adjunto: {exc}"

    expected_ref_norm = _normalize_reference_token(expected_reference)
    found_ref_norm = _extract_reference_from_text(first_page_text)
    has_qr = _has_qr_signal_in_last_page(last_page_text)

    if not found_ref_norm and not has_qr:
        return False, "Se subió el documento equivocado."
    if not found_ref_norm:
        return False, "Se subió el documento equivocado, no se encontró la referencia."
    if expected_ref_norm and found_ref_norm != expected_ref_norm:
        return False, "Subió el documento equivocado, la referencia no concuerda."
    if not has_qr:
        return False, "El documento no está firmado."
    flow_matches, flow_message = _validate_flow_matches_pdf(first_page_text, expected_flow_df)
    if not flow_matches:
        return False, flow_message
    return True, ""

def _complete_drive_oauth_with_code(code: str):
    cfg = st.session_state.get("drive_auth_client_config")
    if not cfg:
        raise RuntimeError("No existe configuración OAuth en sesión. Inicia de nuevo.")
    flow = Flow.from_client_config(cfg, scopes=GOOGLE_DRIVE_UPLOAD_SCOPES)
    flow.redirect_uri = _get_oauth_redirect_uri()
    code_verifier = st.session_state.get("drive_oauth_code_verifier")
    if code_verifier:
        flow.code_verifier = code_verifier
    flow.fetch_token(code=code)
    creds = flow.credentials
    st.session_state.drive_user_token = json.loads(creds.to_json())
    st.session_state.drive_auth_in_progress = False

def _find_col(df: pd.DataFrame, candidates):
    cols = {_norm(c): c for c in df.columns}
    for cand in candidates:
        if _norm(cand) in cols:
            return cols[_norm(cand)]
    return None

####################################################
###################################################

def _find_col_contains(df: pd.DataFrame, required_terms: list[str]):
    normalized_terms = [_norm(term) for term in required_terms]
    for col in df.columns:
        normalized_col = _norm(col)
        if all(term in normalized_col for term in normalized_terms):
            return col
    return None
#####################################################
####################################################

@st.cache_data(ttl=900, show_spinner=False)
def _read_file(file):
    if file.name.lower().endswith(".csv"):
        return pd.read_csv(file)
    return pd.read_excel(file, engine="openpyxl")

@st.cache_data(show_spinner=False)
def _normalize_numeric(df, cols):
    df2 = df.copy()
    for c in cols:
        df2[c] = pd.to_numeric(df2[c].astype(str).str.replace(",", ""), errors="coerce")
    return df2

@st.cache_data(show_spinner=False)
def _map_columns(columns_list: tuple[str, ...]):
    dummy_df = pd.DataFrame(columns=list(columns_list))
    col_ref   = _find_col(dummy_df, ["Referencia"])
    col_id    = _find_col(dummy_df, ["Id deuda","id deuda","id_deuda"])
    col_banco = _find_col(dummy_df, ["Banco"])

    # ✅ MEJORADO: acepta Deuda Resuelve o D_BRAVO (cualquier escritura)
    col_deu   = _find_col(dummy_df, [
        "Deuda Resuelve", "deuda resuelve", "deuda_resuelve", "deuda-resuelve",
        "D_BRAVO", "d_bravo", "d bravo", "d-bravo", "dbravo"
    ])

    col_apar  = _find_col(dummy_df, ["Apartado Mensual","apartado mensual"])
    col_com   = _find_col(dummy_df, ["Comisión Mensual","comision mensual","comisión mensual"])
    col_saldo = _find_col(dummy_df, ["Saldo","Ahorro"])
    col_ce    = _find_col(dummy_df, ["CE"])
    return col_ref, col_id, col_banco, col_deu, col_apar, col_com, col_saldo, col_ce

# ---- Helpers para inputs en pesos con separador de miles (solo otros campos) ----
def _format_pesos(value) -> str:
    try:
        v = float(value)
    except Exception:
        return ""
    if np.isnan(v):
        return ""
    # Formato colombiano: punto miles, coma decimales (máx 2),
    # ocultando decimales cuando sean .00 por defecto.
    return _format_number_es(v, max_decimals=2, trim_trailing=True)

def pesos_input(label: str, key: str, help: str | None = None, disabled: bool = False):
    """
    Input de texto para pesos colombianos (para Deuda, Apartado, etc.).
    """
    raw_val = st.session_state.get(key, 0.0)
    try:
        base_val = float(raw_val or 0.0)
    except Exception:
        base_val = 0.0

    display_key = f"{key}_display"
    synced_key = f"{display_key}_synced"
    if display_key not in st.session_state:
        st.session_state[display_key] = _format_pesos(base_val)
        st.session_state[synced_key] = base_val
    elif not disabled and float(st.session_state.get(synced_key, base_val)) != float(base_val):
        st.session_state[display_key] = _format_pesos(base_val)
        st.session_state[synced_key] = base_val

    txt = st.text_input(label, key=display_key, help=help, disabled=disabled)

    if str(txt).strip() == "":
        new_val = 0.0
    else:
        new_val = _parse_amount_input(txt, max_decimals=2)

    if new_val < 0:
        new_val = 0.0

    st.session_state[key] = new_val
    st.session_state[synced_key] = new_val
    return new_val

# -------- helpers de CE --------
def _prefill_ce():
    """
    Comisión de éxito automática mientras no esté override:
    (Deuda Resuelve - PAGO BANCO) * 1.19 * CE base
    """
    if st.session_state.get("comision_exito_overridden", False):
        return
    deuda_res = float(st.session_state.get("deuda_res_edit", 0.0) or 0.0)
    pago_bco  = float(st.session_state.get("pago_banco", 0.0) or 0.0)
    ce_base   = float(st.session_state.get("ce_base", 0.0) or 0.0)
    ce = max(0.0, (deuda_res - pago_bco) * 1.19 * ce_base)
    st.session_state.comision_exito = ce
    st.session_state.comision_exito_auto = ce  # guardamos referencia del valor "auto"
    
# =================== 1) Cargar base ===================
st.markdown("### 1) Base `cartera_asignada_filtrada`")

DEBUG = False  # ← pon True solo cuando quieras ver debug

df_base = load_repo_base(_data_version())
src_badge = None

# 🔎 DEBUG 1: ¿qué archivo se está usando?
if DEBUG:
    if DATA_PARQUET.exists():
        st.info("📌 Leyendo PARQUET: data/cartera_asignada_filtrada.parquet")
    elif DATA_CSV.exists():
        st.info("📌 Leyendo CSV: data/cartera_asignada_filtrada.csv")
    else:
        st.warning("📌 No hay base en data/, usando subida manual")

# 🔎 DEBUG 2: ver columnas EXACTAS (repr)
if DEBUG and df_base is not None:
    st.write("🧾 Columnas detectadas (repr):")
    st.write([repr(c) for c in df_base.columns])

# 🧹 LIMPIEZA FUERTE de nombres de columnas (quita caracteres invisibles)
def _clean_colname(c):
    c = str(c)
    c = c.replace("\ufeff", "")   # BOM
    c = c.replace("\u200b", "")   # zero-width space
    c = c.replace("\xa0", " ")    # NBSP
    c = c.strip()
    c = re.sub(r"\s+", " ", c)    # colapsa espacios múltiples
    return c

if df_base is not None:
    df_base.columns = [_clean_colname(c) for c in df_base.columns]

# 🔎 DEBUG 3: columnas DESPUÉS de limpiar
if DEBUG and df_base is not None:
    st.write("🧾 Columnas limpiadas:")
    st.write(list(df_base.columns))

###########################################################################################################
if df_base is not None:
    src_badge = "📦 Fuente: data/ (workflow semanal)"
    st.success("✅ Cargada automáticamente desde el repo.")
else:
    src_badge = "📤 Fuente: subida manual"
    st.info("No encontré la base en `data/`. Sube un CSV/XLSX como respaldo.")
    up = st.file_uploader("📂 Subir `cartera_asignada_filtrada`", type=["csv", "xlsx"])
    if not up:
        st.stop()
    try:
        df_base = _read_file(up)
    except Exception as e:
        st.error(f"No pude leer el archivo: {e}")
        st.stop()

st.caption(src_badge)
# Mapear columnas obligatorias
colnames_tuple = tuple(map(str, df_base.columns))
col_ref, col_id, col_banco, col_deu, col_apar, col_com, col_saldo, col_ce = _map_columns(colnames_tuple)

needed = {"Referencia": col_ref, "Id deuda": col_id, "Banco": col_banco,
          "Deuda Resuelve": col_deu, "Apartado Mensual": col_apar,
          "Comisión Mensual": col_com, "Saldo/Ahorro": col_saldo, "CE": col_ce}
faltan = [k for k, v in needed.items() if v is None]
if faltan:
    st.error("Faltan columnas requeridas: " + ", ".join(faltan))
    st.stop()

df_base = _normalize_numeric(df_base, [col_deu, col_apar, col_com, col_saldo, col_ce])
st.success(f"✅ Base lista • filas: {len(df_base):,}")

# =================== 2) Referencia → seleccionar id(s) ===================
st.markdown("### 2) Referencia → seleccionar **Id deuda** (uno o varios)")
ref_input = st.text_input("🔎 Escribe la **Referencia** (exacta como aparece en la base)")
if not ref_input:
    st.stop()

df_ref = df_base[df_base[col_ref].astype(str) == str(ref_input)]
if df_ref.empty:
    st.warning("No encontramos esa referencia en la base.")
    st.stop()

st.subheader("Resultados (elige Id deuda)")
df_preview = (
    df_ref[[col_id, col_banco, col_deu]]
    .rename(columns={col_id: "Id deuda", col_banco: "Banco", col_deu: "Deuda"})
    .copy()
)
df_preview["Id deuda"] = df_preview["Id deuda"].astype(str)
df_preview["Banco"] = df_preview["Banco"].astype(str)
df_preview["Deuda"] = pd.to_numeric(df_preview["Deuda"], errors="coerce").fillna(0.0)

selector_key = f"selector_ids_{ref_input}"
default_ids = set(df_preview["Id deuda"].head(1).tolist())
if selector_key not in st.session_state:
    st.session_state[selector_key] = default_ids

df_selector = df_preview.copy()
df_selector["Seleccionar"] = df_selector["Id deuda"].isin(st.session_state[selector_key])
df_selector["Deuda"] = df_selector["Deuda"].map(_format_currency0)

edited_selector = st.data_editor(
    df_selector[["Seleccionar", "Id deuda", "Banco", "Deuda"]],
    hide_index=True,
    use_container_width=True,
    disabled=["Id deuda", "Banco", "Deuda"],
    column_config={
        "Seleccionar": st.column_config.CheckboxColumn(
            "✅ Seleccionar",
            help="Marca los Id deuda que quieres incluir en el cálculo.",
            default=False,
        )
    },
    key=f"{selector_key}_editor",
)

ids_sel = edited_selector.loc[edited_selector["Seleccionar"], "Id deuda"].astype(str).tolist()
st.session_state[selector_key] = set(ids_sel)
if not ids_sel:
    st.info("Selecciona al menos un Id deuda para continuar.")
    st.stop()

if "modo_sin_portafolio" not in st.session_state:
    st.session_state.modo_sin_portafolio = False

if len(ids_sel) <= 1:
    st.session_state.modo_sin_portafolio = False
else:
    st.toggle(
        "Liquidar sin portafolio",
        key="modo_sin_portafolio",
        help="Activa un flujo separado por deuda (una tabla por Id deuda).",
    )

sel = df_ref[df_ref[col_id].astype(str).isin(ids_sel)].copy()
col_tipo_liquidacion = _find_col(sel, ["Tipo de Liquidacion", "Tipo Liquidacion", "Tipo de liquidación"]) or _find_col_contains(sel, ["tipo", "liquid"])
tipo_liquidacion_val = _join_unique_values(sel[col_tipo_liquidacion].tolist()) if col_tipo_liquidacion else ""
bancos_sel_text = _join_unique_values(sel[col_banco].astype(str).tolist())

# =================== 3) Valores base (reactivo) ===================
st.markdown("### 3) Valores base (puedes editarlos)")

fila_primera = sel.iloc[0]
deuda_res_total_def = float(sel[col_deu].sum(skipna=True))
apartado_base_def   = float(fila_primera[col_apar]) if pd.notna(fila_primera[col_apar]) else 0.0
comision_m_base_def = float(fila_primera[col_com]) if pd.notna(fila_primera[col_com]) else 0.0
saldo_base_def      = float(fila_primera[col_saldo]) if pd.notna(fila_primera[col_saldo]) else 0.0
ce_base_def         = float(fila_primera[col_ce]) if pd.notna(fila_primera[col_ce]) else 0.0

sig_sel = (str(ref_input), tuple(sorted(map(str, ids_sel))))
if st.session_state.get("sig_sel") != sig_sel:
    st.session_state.sig_sel = sig_sel

    st.session_state.deuda_res_edit = deuda_res_total_def
    st.session_state.comision_m_edit = comision_m_base_def
    st.session_state.apartado_edit   = apartado_base_def
    st.session_state.saldo_edit      = saldo_base_def
    st.session_state.ce_base         = ce_base_def

    st.session_state.pago_banco      = 0.0
    st.session_state.n_pab           = 1
    st.session_state.primer_pago_banco = 0.0

    # Flags CE
    st.session_state.comision_exito_overridden = False

    # Comisión de éxito inicial (PAGO BANCO = 0)
    ce_ini = max(0.0, (deuda_res_total_def - 0.0) * 1.19 * ce_base_def)
    st.session_state.comision_exito = ce_ini
    st.session_state.comision_exito_auto = ce_ini

    st.session_state.ce_inicial_val  = 0.0

    st.session_state._last_pab = st.session_state.pago_banco
    st.session_state._last_n   = st.session_state.n_pab

    st.rerun()

col1, col2, col3, col4 = st.columns(4)
with col1:
    pesos_input("💰 Deuda Resuelve", key="deuda_res_edit")
with col2:
    pesos_input("🎯 Comisión Mensual", key="comision_m_edit")
with col3:
    pesos_input("📆 Apartado Mensual", key="apartado_edit")
with col4:
    pesos_input("💼 Saldo (Ahorro)", key="saldo_edit")

# 3.4 Saldo Neto y Depósito
saldo_neto = 0.0
if pd.notna(st.session_state.saldo_edit) and st.session_state.saldo_edit > 0:
    saldo_neto = float(st.session_state.saldo_edit) - (float(st.session_state.saldo_edit) - 25000.0) * 0.004
    saldo_neto = max(0.0, saldo_neto)
saldo_neto_disp = float(np.round(saldo_neto, 0))

col5, col6 = st.columns(2)
with col5:
    st.number_input(
        "🧾 Saldo Neto",
        value=saldo_neto_disp,
        step=1000.0,
        min_value=0.0,
        format="%.0f",
        disabled=True,
        help="Saldo − (Saldo − 25.000) × 0.004 (solo si Saldo > 0)"
    )
with col6:
    pesos_input("💵 Depósito", key="deposito_edit",
                help="Monto extra aportado al inicio; por defecto 0")

# =================== 4) PAGO BANCO y derivados ===================
st.markdown("### 4) PAGO BANCO y parámetros derivados")

c1, c2, c3 = st.columns([1,1,1])
with c1:
    pesos_input("🏦 PAGO BANCO", key="pago_banco")
with c2:
    descuento = None
    if st.session_state.deuda_res_edit and st.session_state.deuda_res_edit > 0:
        descuento = max(0.0, 1.0 - (st.session_state.pago_banco / st.session_state.deuda_res_edit)) * 100.0
    st.text_input(
        "📉 DESCUENTO (%)",
        value=(f"{descuento:.2f} %" if descuento is not None else ""),
        disabled=True
    )
with c3:
    st.number_input(
        "🧮 N PaB",
        min_value=1,
        step=1,
        value=int(st.session_state.n_pab),
        key="n_pab"
    )

# --- Lógica: si cambia N PaB, recalculamos Primer PAGO BANCO ---
pago_banco = float(st.session_state.get("pago_banco", 0.0) or 0.0)
n_pab = int(st.session_state.get("n_pab", 1) or 1)

prev_n_pab = st.session_state.get("_prev_n_pab_for_primer", n_pab)
if n_pab != prev_n_pab:
    # Cambió el N PaB → recalculamos primer pago
    if n_pab > 1:
        if pago_banco > 0:
            st.session_state.primer_pago_banco = pago_banco / n_pab
        else:
            st.session_state.primer_pago_banco = 0.0
    else:
        # Si vuelve a 1, todo el PAGO BANCO va al primer pago
        st.session_state.primer_pago_banco = pago_banco
st.session_state._prev_n_pab_for_primer = n_pab
# --------------------------------------------------------

# Campo adicional: Primer PAGO BANCO solo si N PaB > 1
if n_pab > 1:
    pago_banco_actual = float(st.session_state.pago_banco or 0.0)

    # Aseguramos que no supere el total ni sea negativo
    st.session_state.primer_pago_banco = min(
        max(float(st.session_state.get("primer_pago_banco", 0.0) or 0.0), 0.0),
        pago_banco_actual
    )

    pesos_input(
        "💳 Primer PAGO BANCO",
        key="primer_pago_banco",
        help="Monto del primer pago al banco (el resto se reparte en los siguientes PaB)."
    )
else:
    # Si solo hay un PaB, el primer pago es todo el PAGO BANCO
    st.session_state.primer_pago_banco = float(st.session_state.pago_banco or 0.0)

# Detectar cambios en PAGO BANCO / N PaB para recalcular CE
if (st.session_state._last_pab != st.session_state.pago_banco) or (st.session_state._last_n != st.session_state.n_pab):
    st.session_state._last_pab = st.session_state.pago_banco
    st.session_state._last_n   = st.session_state.n_pab
    _prefill_ce()

# Comisión de éxito (editable) y CE inicial
c4, c5 = st.columns(2)
with c4:
    # Valor automático de referencia
    auto_ce = float(st.session_state.get("comision_exito_auto", st.session_state.get("comision_exito", 0.0)) or 0.0)
    new_ce = pesos_input(
        "🏁 Comisión de éxito (editable)",
        key="comision_exito",
        help="Por defecto se calcula con la fórmula, pero puedes ajustarla manualmente."
    )

    # Si el valor actual se separa del valor "auto", marcamos override
    if not st.session_state.get("comision_exito_overridden", False):
        if abs(new_ce - auto_ce) > 0.5:
            st.session_state.comision_exito_overridden = True

with c5:
    pesos_input("🧪 CE inicial", key="ce_inicial_val")

# Avance CE inicial vs Comisión de éxito
st.markdown("#### Avance de CE inicial sobre la Comisión de éxito")
ce_inicial = float(st.session_state.get("ce_inicial_val", 0.0) or 0.0)
base = float(st.session_state.get("comision_exito", 0.0) or 0.0)
if ce_inicial <= 0:
    st.info("Escribe un valor en **CE inicial** para ver el porcentaje.")
elif base <= 0:
    st.warning("La **Comisión de éxito** debe ser mayor a 0 para calcular el porcentaje.")
else:
    porcentaje = (ce_inicial / base) * 100.0
    porcentaje_capped = max(0.0, min(porcentaje, 100.0))
    st.progress(int(round(porcentaje_capped)))
    st.caption(f"CE inicial: {ce_inicial:,.0f}  |  Comisión de éxito: {base:,.0f}  →  **{porcentaje:,.2f}%**")

# =================== 6) Validación y KPIs (sin tabla) ===================
st.markdown("### 6) Validación y KPIs")

pago_banco        = float(st.session_state.get("pago_banco", 0.0) or 0.0)
n_pab             = int(st.session_state.get("n_pab", 1) or 1)
comision_mensual  = float(st.session_state.get("comision_m_edit", 0.0) or 0.0)
apartado_mensual  = float(st.session_state.get("apartado_edit", 0.0) or 0.0)
comision_exito    = float(st.session_state.get("comision_exito", 0.0) or 0.0)
ce_inicial        = float(st.session_state.get("ce_inicial_val", 0.0) or 0.0)

plazo = st.number_input("📅 PLAZO (meses) (lo ingresas tú)", min_value=1, step=1, value=1)

# Primer PAGO BANCO: si hay más de un PaB, usamos el input; si no, todo el pago
primer_pago_banco = float(
    st.session_state.get(
        "primer_pago_banco",
        pago_banco if n_pab == 1 else (pago_banco / n_pab if n_pab > 0 else 0.0)
    )
)
primer_pago_banco = min(max(primer_pago_banco, 0.0), pago_banco)
resto_pago_banco = max(0.0, pago_banco - primer_pago_banco)

pct_primer_pago = (ce_inicial / comision_exito) if comision_exito > 0 else np.nan

if (plazo - 1) > 0 and apartado_mensual > 0:
    numerador = (comision_exito + resto_pago_banco - ce_inicial + comision_mensual)
    cuota_apartado = (numerador / (plazo - 1)) / apartado_mensual
else:
    cuota_apartado = np.nan

cronograma_df, cronograma_meta = construir_cronograma_pagos(
    fecha_inicial=date.today(),
    plazo=int(plazo),
    n_pab=n_pab,
    pago_banco_total=pago_banco,
    primer_pago_banco=primer_pago_banco,
    comision_total=comision_exito,
    comision_inicial=ce_inicial,
)

c1, c2, c3, c4 = st.columns(4)
with c1:
    st.number_input("🏁 Comisión de éxito total", value=float(comision_exito), step=0.0, format="%.0f", disabled=True)
with c2:
    st.number_input("📅 PLAZO (meses)", value=float(plazo), step=1.0, format="%.0f", disabled=True)
with c3:
    st.text_input(
        "% Primer Pago (CE inicial / CE)",
        value=("—" if np.isnan(pct_primer_pago) else f"{pct_primer_pago:.2f}"),
        disabled=True
    )
with c4:
    st.text_input(
        "Cuota/Apartado",
        value=("—" if np.isnan(cuota_apartado) else f"{cuota_apartado:.4f}"),
        disabled=True
    )

if st.session_state.get("comision_exito_overridden", False):
    st.caption("🔒 Comisión de éxito fijada manualmente: no se recalcula con cambios en PAGO BANCO/N PaB.")
############################################################################################################################################################################
st.markdown("### 6.1) Flujo sugerido de pagos")

fecha_cfg_1, fecha_cfg_2, fecha_cfg_3, fecha_cfg_4, fecha_cfg_5 = st.columns([1.2, 1, 1.2, 1, 1])
with fecha_cfg_1:
    modo_fecha_banco = st.radio("Fechas banco", options=["Fin de mes", "Día fijo"], horizontal=True, key="modo_fecha_banco")
with fecha_cfg_2:
    dia_pago_banco = None
    if modo_fecha_banco == "Día fijo":
        dia_pago_banco = int(st.number_input("Día banco", min_value=1, max_value=31, value=int(st.session_state.get("dia_pago_banco", 15)), step=1, key="dia_pago_banco"))
    else:
        st.caption("Banco: fin de mes.")
with fecha_cfg_3:
    modo_fecha_comision = st.radio("Fechas comisión", options=["Fin de mes", "Día fijo"], horizontal=True, key="modo_fecha_comision")
with fecha_cfg_4:
    dia_pago_comision = None
    if modo_fecha_comision == "Día fijo":
        dia_pago_comision = int(st.number_input("Día comisión", min_value=1, max_value=31, value=int(st.session_state.get("dia_pago_comision", 15)), step=1, key="dia_pago_comision"))
    else:
        st.caption("Comisión: fin de mes.")
with fecha_cfg_5:
    if st.button("Restablecer cronograma", use_container_width=True):
        st.session_state.pop("cronograma_editor", None)
        st.session_state.pop("cronograma_overrides", None)
        st.rerun()

cronograma_df, cronograma_meta = construir_cronograma_pagos(
    fecha_inicial=date.today(),
    plazo=int(plazo),
    n_pab=n_pab,
    pago_banco_total=pago_banco,
    primer_pago_banco=primer_pago_banco,
    comision_total=comision_exito,
    comision_inicial=ce_inicial,
    dia_pago_banco=dia_pago_banco,
    dia_pago_comision=dia_pago_comision,
)

totales_por_tipo = {"banco": float(pago_banco), "comision": float(comision_exito)}
cronogramas_individuales: dict[str, pd.DataFrame] = {}
expected_flow_for_pdf_validation = None
if st.session_state.get("modo_sin_portafolio", False) and len(ids_sel) > 1:
    st.info("Modo sin portafolio activo: se crea un flujo independiente por cada Id deuda seleccionado.")
    per_debt_key = f"sin_portafolio_config_{ref_input}"
    deuda_detalle = (
        sel[[col_id, col_deu]]
        .rename(columns={col_id: "Id deuda", col_deu: "Deuda Resuelve"})
        .copy()
    )
    deuda_detalle["Id deuda"] = deuda_detalle["Id deuda"].astype(str)
    deuda_detalle["Deuda Resuelve"] = pd.to_numeric(deuda_detalle["Deuda Resuelve"], errors="coerce").fillna(0.0)
    deuda_detalle = deuda_detalle.sort_values(["Deuda Resuelve", "Id deuda"], ascending=[False, True]).reset_index(drop=True)
    pesos = deuda_detalle["Deuda Resuelve"].tolist()
    pago_banco_deuda = _split_total_by_weights(pago_banco, pesos, digits=2)
    ce_base_mode = float(st.session_state.get("ce_base", 0.0) or 0.0)
    comision_deuda = [
        max((float(deuda_detalle.iloc[i]["Deuda Resuelve"]) - float(pago_banco_deuda[i])) * 1.19 * ce_base_mode, 0.0)
        for i in range(len(deuda_detalle))
    ]
    comision_deuda = _sum_rounded_parts(comision_deuda, digits=2)
    primer_banco_deuda = _split_total_by_weights(primer_pago_banco, pesos, digits=2)
    primer_comision_deuda = _split_total_by_weights(ce_inicial, comision_deuda, digits=2)
    n_pab_deuda = _split_integer_equitable(n_pab, len(deuda_detalle))

    default_conf = pd.DataFrame({
        "Id deuda": deuda_detalle["Id deuda"],
        "Deuda Resuelve": deuda_detalle["Deuda Resuelve"],
        "Pago Banco Total": pago_banco_deuda,
        "Comisión Total": comision_deuda,
        "Primer Pago Banco": primer_banco_deuda,
        "CE Inicial": primer_comision_deuda,
        "N PaB": n_pab_deuda,
    })
    config_view = default_conf[["Id deuda", "Deuda Resuelve", "Pago Banco Total", "Comisión Total", "Primer Pago Banco", "CE Inicial", "N PaB"]].copy()
    config_view["Deuda Resuelve"] = config_view["Deuda Resuelve"].map(_format_currency0)
    config_view["Pago Banco Total"] = config_view["Pago Banco Total"].map(_format_currency0)
    config_view["Comisión Total"] = config_view["Comisión Total"].map(_format_currency0)
    config_view["Primer Pago Banco"] = config_view["Primer Pago Banco"].map(_format_currency0)
    config_view["CE Inicial"] = config_view["CE Inicial"].map(_format_currency0)
    st.caption("Los totales por deuda se asignan por proporción. Puedes editar Pago Banco Total, Primer Pago Banco y CE Inicial.")
    edited_conf = st.data_editor(
        config_view,
        key=per_debt_key,
        use_container_width=True,
        hide_index=True,
        num_rows="fixed",
        disabled=["Id deuda", "Deuda Resuelve", "Comisión Total"],
    )

    pago_banco_edit = [_parse_amount_input(x) for x in edited_conf["Pago Banco Total"].tolist()]
    default_pagos = [float(x) for x in default_conf["Pago Banco Total"].tolist()]
    pago_edit_flags = [abs(float(pago_banco_edit[i]) - default_pagos[i]) > 1.0 for i in range(len(default_pagos))]
    idx_editados = [i for i, flag in enumerate(pago_edit_flags) if flag]
    idx_no_editados = [i for i, flag in enumerate(pago_edit_flags) if not flag]

    if idx_editados:
        suma_editados = sum(max(float(pago_banco_edit[i]), 0.0) for i in idx_editados)
        if suma_editados > float(pago_banco) + 0.01:
            st.warning("La suma de Pago Banco Total editada no puede superar el PAGO BANCO global. Se ajustó automáticamente.")
            scaled_vals = _split_total_by_weights(float(pago_banco), [pago_banco_edit[i] for i in idx_editados], digits=2)
            for idx_loc, valor in zip(idx_editados, scaled_vals):
                pago_banco_edit[idx_loc] = valor
            suma_editados = sum(max(float(pago_banco_edit[i]), 0.0) for i in idx_editados)

        restante = max(float(pago_banco) - suma_editados, 0.0)
        if idx_no_editados:
            pesos_restantes = [float(deuda_detalle.iloc[i]["Deuda Resuelve"]) for i in idx_no_editados]
            partes_restantes = _split_total_by_weights(restante, pesos_restantes, digits=2)
            for idx_loc, valor in zip(idx_no_editados, partes_restantes):
                pago_banco_edit[idx_loc] = valor
        else:
            pago_banco_edit[-1] = round(max(float(pago_banco) - sum(pago_banco_edit[:-1]), 0.0), 2)
    else:
        pago_banco_edit = default_pagos

    max_comision_global = float(st.session_state.get("comision_exito", comision_exito) or 0.0)
    comision_deuda_calc = [
        max((float(deuda_detalle.iloc[i]["Deuda Resuelve"]) - float(pago_banco_edit[i])) * 1.19 * ce_base_mode, 0.0)
        for i in range(len(deuda_detalle))
    ]
    comision_deuda_calc = _sum_rounded_parts(comision_deuda_calc, digits=2)
    total_comision_calc = float(sum(comision_deuda_calc))
    if total_comision_calc > max_comision_global + 0.01:
        st.warning("La Comisión de éxito recalculada superó el tope global; se ajustó proporcionalmente.")
        comision_deuda_calc = _split_total_by_weights(max_comision_global, comision_deuda_calc, digits=2)
    ce_inicial_default = _split_total_by_weights(ce_inicial, comision_deuda_calc, digits=2)
    default_conf["Pago Banco Total"] = pago_banco_edit
    default_conf["Comisión Total"] = comision_deuda_calc
    default_conf["CE Inicial"] = ce_inicial_default
    comision_exito = float(sum(comision_deuda_calc))
    st.caption(f"Comisión de éxito recalculada (sin portafolio): {_format_currency0(comision_exito)}")

    primeres_banco_edit = [_parse_amount_input(x) for x in edited_conf["Primer Pago Banco"].tolist()]
    primeras_com_edit = [_parse_amount_input(x) for x in edited_conf["CE Inicial"].tolist()]
    ce_default_vals = [float(x) for x in ce_inicial_default]
    ce_edit_flags = [abs(float(primeras_com_edit[i]) - ce_default_vals[i]) > 1.0 for i in range(len(ce_default_vals))]
    ce_idx_editados = [i for i, flag in enumerate(ce_edit_flags) if flag]
    ce_idx_no_editados = [i for i, flag in enumerate(ce_edit_flags) if not flag]
    if ce_idx_editados:
        suma_ce_editados = sum(max(float(primeras_com_edit[i]), 0.0) for i in ce_idx_editados)
        if suma_ce_editados > float(ce_inicial) + 0.01:
            st.warning("La suma de CE Inicial por deuda no puede superar el CE Inicial global. Se ajustó automáticamente.")
            ce_scaled = _split_total_by_weights(float(ce_inicial), [primeras_com_edit[i] for i in ce_idx_editados], digits=2)
            for idx_loc, valor in zip(ce_idx_editados, ce_scaled):
                primeras_com_edit[idx_loc] = valor
            suma_ce_editados = sum(max(float(primeras_com_edit[i]), 0.0) for i in ce_idx_editados)
        ce_restante = max(float(ce_inicial) - suma_ce_editados, 0.0)
        if ce_idx_no_editados:
            ce_pesos_rest = [comision_deuda_calc[i] for i in ce_idx_no_editados]
            ce_partes = _split_total_by_weights(ce_restante, ce_pesos_rest, digits=2)
            for idx_loc, valor in zip(ce_idx_no_editados, ce_partes):
                primeras_com_edit[idx_loc] = valor
        else:
            primeras_com_edit[-1] = round(max(float(ce_inicial) - sum(primeras_com_edit[:-1]), 0.0), 2)
    else:
        primeras_com_edit = ce_default_vals
    n_pab_edit = [max(int(_parse_amount_input(x, max_decimals=0)), 1) for x in edited_conf["N PaB"].tolist()]
    n_pab_diff = int(n_pab) - sum(n_pab_edit)
    if n_pab_edit:
        n_pab_edit[-1] = max(1, n_pab_edit[-1] + n_pab_diff)
    if sum(n_pab_edit) != int(n_pab):
        st.warning("No fue posible cuadrar N PaB por deuda con el total; se aplicó ajuste automático por descarte.")
        n_pab_edit = _split_integer_equitable(n_pab, len(deuda_detalle))
    limite_banco = float(pago_banco)
    limite_comision = float(ce_inicial)
    if sum(primeres_banco_edit) > limite_banco + 0.01:
        st.warning("La suma de Primer Pago Banco por deuda no puede superar el PAGO BANCO total. Se ajustó automáticamente.")
        primeres_banco_edit = _split_total_by_weights(min(limite_banco, sum(primeres_banco_edit)), pesos, digits=2)
    if sum(primeras_com_edit) > limite_comision + 0.01:
        st.warning("La suma de CE Inicial por deuda no puede superar el CE Inicial global. Se ajustó automáticamente.")
        primeras_com_edit = _split_total_by_weights(min(limite_comision, sum(primeras_com_edit)), pesos, digits=2)

    cronos = []
    for i, row in default_conf.iterrows():
        deuda_id = str(row["Id deuda"])
        primer_banco_i = min(max(primeres_banco_edit[i], 0.0), float(row["Pago Banco Total"]))
        primera_com_i = min(max(primeras_com_edit[i], 0.0), float(row["Comisión Total"]))
        cr_df_i, _ = construir_cronograma_pagos(
            fecha_inicial=date.today(),
            plazo=int(plazo),
            n_pab=int(max(1, n_pab_edit[i])),
            pago_banco_total=float(row["Pago Banco Total"]),
            primer_pago_banco=primer_banco_i,
            comision_total=float(row["Comisión Total"]),
            comision_inicial=primera_com_i,
            dia_pago_banco=dia_pago_banco,
            dia_pago_comision=dia_pago_comision,
        )
        cr_df_i["Id deuda"] = deuda_id
        deuda_overrides_key = f"cronograma_overrides_sin_portafolio_{ref_input}_{deuda_id}"
        deuda_editor_key = f"cronograma_editor_sin_portafolio_{ref_input}_{deuda_id}"
        deuda_overrides = st.session_state.get(deuda_overrides_key, {})
        deuda_editor_state = st.session_state.get(deuda_editor_key, {})
        for row_position_str, cambios in (deuda_editor_state.get("edited_rows", {}) or {}).items():
            try:
                row_position = int(row_position_str)
            except (TypeError, ValueError):
                continue
            visible_i = cr_df_i[cr_df_i["Cantidad"] > 0.005].reset_index(drop=True)
            if row_position < 0 or row_position >= len(visible_i):
                continue
            row_key = str(visible_i.iloc[row_position]["row_key"])
            existing = deuda_overrides.get(row_key, {})
            existing.update(cambios)
            deuda_overrides[row_key] = existing
        st.session_state[deuda_overrides_key] = deuda_overrides

        cr_df_i_editado, advertencias_i = aplicar_overrides_cronograma(
            cronograma_df=cr_df_i,
            overrides_map=deuda_overrides,
            totales_por_tipo={"banco": float(row["Pago Banco Total"]), "comision": float(row["Comisión Total"])},
            fecha_inicial=date.today(),
            dia_pago_banco=dia_pago_banco,
            dia_pago_comision=dia_pago_comision,
            primer_pago_banco_input=primer_banco_i,
            comision_inicial_input=primera_com_i,
            lock_initial_rows=False,
        )
        for advertencia in advertencias_i:
            st.warning(f"Id deuda {deuda_id}: {advertencia}")

        cronogramas_individuales[deuda_id] = cr_df_i_editado.copy()
        if expected_flow_for_pdf_validation is None:
            expected_flow_for_pdf_validation = cr_df_i_editado[cr_df_i_editado["Cantidad"] > 0.005][["Fecha", "Cantidad", "Concepto"]].copy()
        cronos.append(cr_df_i_editado.copy())

        st.markdown(f"#### Flujo Id deuda {deuda_id}")
        table_i = cr_df_i_editado[cr_df_i_editado["Cantidad"] > 0.005][["Fecha", "Cantidad", "Concepto"]].copy()
        if not table_i.empty:
            table_i["Fecha"] = pd.to_datetime(table_i["Fecha"])
            table_i["Cantidad"] = table_i["Cantidad"].round(2).map(_format_currency0)
            table_i.index = range(1, len(table_i) + 1)
            st.data_editor(
                table_i,
                key=deuda_editor_key,
                use_container_width=True,
                hide_index=False,
                num_rows="fixed",
                disabled=["Concepto"],
                column_config={
                    "Fecha": st.column_config.DateColumn("Fecha", format="DD/MM/YYYY"),
                    "Cantidad": st.column_config.TextColumn("Cantidad"),
                    "Concepto": st.column_config.TextColumn("Concepto", disabled=True),
                },
            )

    if cronos:
        cronograma_editado = pd.concat(cronos, ignore_index=True)
        cronograma_editado = (
            cronograma_editado
            .groupby(["Fecha", "Concepto"], as_index=False)["Cantidad"]
            .sum()
            .sort_values("Fecha")
            .reset_index(drop=True)
        )
    else:
        cronograma_editado = pd.DataFrame(columns=["Fecha", "Cantidad", "Concepto"])
else:
    cronograma_overrides = st.session_state.get("cronograma_overrides", {})
    cronograma_editor_state = st.session_state.get("cronograma_editor", {})

    cronograma_base_editado, _ = aplicar_overrides_cronograma(
        cronograma_df=cronograma_df,
        overrides_map=cronograma_overrides,
        totales_por_tipo=totales_por_tipo,
        fecha_inicial=date.today(),
        dia_pago_banco=dia_pago_banco,
        dia_pago_comision=dia_pago_comision,
        primer_pago_banco_input=primer_pago_banco,
        comision_inicial_input=ce_inicial,
    )

    cronograma_base_visible = cronograma_base_editado[cronograma_base_editado["Cantidad"] > 0.005].reset_index(drop=True)
    cronograma_locked_rows_changed = False
    for row_position_str, cambios in (cronograma_editor_state.get("edited_rows", {}) or {}).items():
        try:
            row_position = int(row_position_str)
        except (TypeError, ValueError):
            continue
        if row_position < 0 or row_position >= len(cronograma_base_visible):
            continue
        row = cronograma_base_visible.iloc[row_position]
        if int(row["months_ahead"]) == 0:
            cronograma_locked_rows_changed = True
            continue
        row_key = str(row["row_key"])
        existing = cronograma_overrides.get(row_key, {})
        existing.update(cambios)
        cronograma_overrides[row_key] = existing
    st.session_state["cronograma_overrides"] = cronograma_overrides

    if cronograma_locked_rows_changed:
        st.session_state.pop("cronograma_editor", None)
        st.info("Las filas 1 y 2 están bloqueadas y no se pueden editar.")
        st.rerun()

    cronograma_editado, advertencias_cronograma = aplicar_overrides_cronograma(
        cronograma_df=cronograma_df,
        overrides_map=cronograma_overrides,
        totales_por_tipo=totales_por_tipo,
        fecha_inicial=date.today(),
        dia_pago_banco=dia_pago_banco,
        dia_pago_comision=dia_pago_comision,
        primer_pago_banco_input=primer_pago_banco,
        comision_inicial_input=ce_inicial,
    )

    for advertencia in advertencias_cronograma:
        st.warning(advertencia)

    cronograma_visible = cronograma_editado[cronograma_editado["Cantidad"] > 0.005].copy()
    expected_flow_for_pdf_validation = cronograma_visible[["Fecha", "Cantidad", "Concepto"]].copy() if not cronograma_visible.empty else None
    if not cronograma_visible.empty:
        cronograma_view = cronograma_visible[["Fecha", "Cantidad", "Concepto"]].copy()
        cronograma_view["Fecha"] = pd.to_datetime(cronograma_view["Fecha"])
        cronograma_view["Cantidad"] = (
            pd.to_numeric(cronograma_view["Cantidad"], errors="coerce")
            .fillna(0.0)
            .round(2)
            .map(_format_currency0)
        )
        cronograma_view.index = range(1, len(cronograma_view) + 1)
        st.caption("Sugerencia: banco y comisión van en meses diferentes, pero si mueves una comisión al mismo mes del banco se respeta y las demás comisiones siguen ocupando los meses restantes sin dejar huecos.")
        st.caption("Filas 1 y 2 bloqueadas (no editables).")
        st.data_editor(
            cronograma_view,
            key="cronograma_editor",
            use_container_width=True,
            num_rows="fixed",
            hide_index=False,
            column_config={
                "Fecha": st.column_config.DateColumn("Fecha", format="DD/MM/YYYY"),
                "Cantidad": st.column_config.TextColumn("Cantidad"),
                "Concepto": st.column_config.TextColumn("Concepto", disabled=True),
            },
            disabled=["Concepto"],
        )
    else:
        st.info("Aún no hay valores suficientes para construir el cronograma.")

st.markdown("### PLAN DE LIQUIDACIÓN ESTRUCTURADA")

plan_df = construir_plan_liquidacion(cronograma_editado, comision_mensual)
plan_overrides = st.session_state.get("plan_liquidacion_overrides", {})
plan_editor_state = st.session_state.get("plan_liquidacion_editor", {})

for row_position_str, cambios in (plan_editor_state.get("edited_rows", {}) or {}).items():
    try:
        row_position = int(row_position_str)
    except (TypeError, ValueError):
        continue
    if row_position < 0 or row_position >= len(plan_df):
        continue
    plan_key = str(plan_df.iloc[row_position]["plan_key"])
    if "Comisión Mensual" in cambios:
        plan_overrides[plan_key] = _parse_amount_input(cambios["Comisión Mensual"])

st.session_state["plan_liquidacion_overrides"] = plan_overrides

if not plan_df.empty:
    plan_df["Comisión Mensual"] = plan_df.apply(
        lambda row: float(plan_overrides.get(str(row["plan_key"]), row["Comisión Mensual"])),
        axis=1,
    )
    plan_df["Apartado Requerido"] = plan_df["Pago a Banco"] + plan_df["Comisión de Éxito"] + plan_df["Comisión Mensual"]

    plan_view = plan_df.copy()
    plan_view["Fecha Límite de Pago"] = pd.to_datetime(plan_view["Fecha Límite de Pago"])
    plan_view["Pago a Banco"] = plan_view["Pago a Banco"].map(_format_currency0)
    plan_view["Comisión de Éxito"] = plan_view["Comisión de Éxito"].map(_format_currency0)
    plan_view["Comisión Mensual"] = plan_view["Comisión Mensual"].map(_format_currency0)
    plan_view["Apartado Requerido"] = plan_view["Apartado Requerido"].map(_format_currency0)
    plan_view = plan_view.drop(columns=["plan_key"])
    plan_view.index = range(1, len(plan_view) + 1)

    st.data_editor(
        plan_view,
        key="plan_liquidacion_editor",
        use_container_width=True,
        num_rows="fixed",
        hide_index=False,
        column_config={
            "Fecha Límite de Pago": st.column_config.DateColumn("Fecha Límite de Pago", format="DD/MM/YYYY"),
            "Pago a Banco": st.column_config.TextColumn("Pago a Banco", disabled=True),
            "Comisión de Éxito": st.column_config.TextColumn("Comisión de Éxito", disabled=True),
            "Comisión Mensual": st.column_config.TextColumn("Comisión Mensual"),
            "Apartado Requerido": st.column_config.TextColumn("Apartado Requerido", disabled=True),
        },
        disabled=["Pago a Banco", "Comisión de Éxito", "Apartado Requerido"],
    )
else:
    st.info("Aún no hay datos suficientes para construir el plan de liquidación.")

#############################################################################################################################################################################

def _build_document_context_inputs(default_context: dict[str, str]) -> dict[str, str]:
    with st.expander("Campos editables del documento", expanded=False):
        st.caption("Antes de descargar el Word, puedes revisar y corregir aquí los datos que se van a escribir en la plantilla.")

        col1, col2 = st.columns(2)
        with col1:
            referencia_doc = st.text_input("Referencia documento", value=str(default_context.get("referencia", "")), key="doc_referencia")
            dia_firma_doc = st.text_input("Día firma", value=str(default_context.get("dia_firma", "")), key="doc_dia_firma", disabled=True)
            mes_firma_doc = st.text_input("Mes firma", value=str(default_context.get("mes_firma", "")), key="doc_mes_firma", disabled=True)
            anio_firma_doc = st.text_input("Año firma", value=str(default_context.get("anio_firma", "")), key="doc_anio_firma", disabled=True)
            entidad_financiera_doc = st.text_input("Entidad financiera", value=str(default_context.get("entidad_financiera", "")), key="doc_entidad_financiera")
            nombre_cliente_doc = st.text_input("Nombre cliente", value=str(default_context.get("nombre_cliente", "")), key="doc_nombre_cliente")
            correo_cliente_doc = st.text_input("Correo cliente", value=str(default_context.get("correo_cliente", "")), key="doc_correo_cliente")
            telefono_cliente_doc = st.text_input("Teléfono cliente", value=str(default_context.get("telefono_cliente", "")), key="doc_telefono_cliente")
        with col2:
            numero_producto_doc = st.text_input("Número producto", value=str(default_context.get("numero_producto", "")), key="doc_numero_producto")
            vehiculo_doc = st.text_input("Vehículo", value=str(default_context.get("vehiculo", "")), key="doc_vehiculo")
            cedula_cliente_doc = st.text_input("Cédula cliente", value=str(default_context.get("cedula_cliente", "")), key="doc_cedula_cliente")
            ciudad_cliente_doc = st.text_input("Ciudad cliente", value=str(default_context.get("ciudad_cliente", "")), key="doc_ciudad_cliente")
            direccion_cliente_doc = st.text_input("Dirección cliente", value=str(default_context.get("direccion_cliente", "")), key="doc_direccion_cliente")
            pago_banco_doc = st.text_input("Pago banco documento", value=str(default_context.get("pago_banco", "")), key="doc_pago_banco", disabled=True)
            comision_total_doc = st.text_input("Comisión total documento", value=str(default_context.get("comision_total", "")), key="doc_comision_total", disabled=True)

    context = default_context.copy()
    context.update({
        "referencia": referencia_doc,
        "dia_firma": dia_firma_doc,
        "mes_firma": mes_firma_doc,
        "anio_firma": anio_firma_doc,
        "entidad_financiera": entidad_financiera_doc,
        "nombre_cliente": nombre_cliente_doc,
        "correo_cliente": correo_cliente_doc,
        "telefono_cliente": telefono_cliente_doc,
        "numero_producto": numero_producto_doc,
        "vehiculo": vehiculo_doc,
        "cedula_cliente": cedula_cliente_doc,
        "ciudad_cliente": ciudad_cliente_doc,
        "direccion_cliente": direccion_cliente_doc,
        "pago_banco": pago_banco_doc,
        "comision_total": comision_total_doc,
        "suma_comisiones": default_context.get("suma_comisiones", comision_total_doc),
        "Suma_comisiones": default_context.get("suma_comisiones", comision_total_doc),
        "suma comisiones": default_context.get("suma_comisiones", comision_total_doc),
        "suma_comisiones_letras": default_context.get("suma_comisiones_letras", ""),
    })
    return context


def _missing_document_fields(context: dict[str, str]) -> list[str]:
    def _is_empty_or_placeholder(value) -> bool:
        text = str(value or "").strip()
        if not text:
            return True
        normalized = re.sub(r"[\s<>]", "", text).lower()
        return normalized in {"na", "n/a", "nan", "none", "null"}
        
    required_labels = {
        "referencia": "Referencia documento",
        "dia_firma": "Día firma",
        "mes_firma": "Mes firma",
        "anio_firma": "Año firma",
        "entidad_financiera": "Entidad financiera",
        "nombre_cliente": "Nombre cliente",
        "correo_cliente": "Correo cliente",
        "telefono_cliente": "Teléfono cliente",
        "numero_producto": "Número producto",
        "vehiculo": "Vehículo",
        "cedula_cliente": "Cédula cliente",
        "ciudad_cliente": "Ciudad cliente",
        "direccion_cliente": "Dirección cliente",
        "pago_banco": "Pago banco documento",
        "comision_total": "Comisión total documento",
    }
    missing = []
    for key, label in required_labels.items():
        if _is_empty_or_placeholder(context.get(key, "")):
            missing.append(label)
    return missing


st.markdown("### 6.2) Exportar documento estructurado")
st.caption(
    "La plantilla Word de `data/` se rellena con las dos tablas visibles en pantalla. "
    "Las celdas exportadas se normalizan en Times New Roman 9.5 para que el estilo interno coincida con el resto del documento."
)

if "doc_graduacion_confirmada" not in st.session_state:
    st.session_state.doc_graduacion_confirmada = False
if "doc_graduacion_pendiente" not in st.session_state:
    st.session_state.doc_graduacion_pendiente = False
if "doc_graduacion_check" not in st.session_state:
    st.session_state.doc_graduacion_check = False

st.checkbox(
    "Graduar cliente",
    key="doc_graduacion_check",
    help="Incluye el punto 6 de la primera página solo después de confirmarlo.",
)

if not st.session_state.doc_graduacion_check:
    st.session_state.doc_graduacion_confirmada = False
    st.session_state.doc_graduacion_pendiente = False
elif not st.session_state.doc_graduacion_confirmada:
    st.session_state.doc_graduacion_pendiente = True

if st.session_state.doc_graduacion_pendiente:
    st.caption("Confirma si el cliente sí se va a graduar.")
    col_confirmar_si, col_confirmar_no = st.columns([1, 1, 4])[:2]
    with col_confirmar_si:
        if st.button("Sí", key="confirmar_graduacion_si"):
            st.session_state.doc_graduacion_confirmada = True
            st.session_state.doc_graduacion_pendiente = False
            st.rerun()
    with col_confirmar_no:
        if st.button("No", key="confirmar_graduacion_no"):
            st.session_state.doc_graduacion_check = False
            st.session_state.doc_graduacion_confirmada = False
            st.session_state.doc_graduacion_pendiente = False
            st.rerun()
elif st.session_state.doc_graduacion_check and st.session_state.doc_graduacion_confirmada:
    st.caption("Graduación confirmada: el Word incluirá el punto 6 en la primera página.")

export_pdf_bytes = None
export_pdf_error = None
if not cronograma_editado.empty and not plan_df.empty:
    try:
        col_nombre_cliente = _find_col(sel, ["Nombre del cliente", "Nombre Cliente", "Nombre"]) or _find_col_contains(sel, ["nombre", "cliente"])
        col_numero_producto = _find_col(sel, ["Número de Crédito", "Numero de Credito", "Número Crédito", "Numero Producto"]) or _find_col_contains(sel, ["numero", "credito"])
        col_vehiculo = _find_col(sel, ["vehiculo", "Vehículo"]) or _find_col_contains(sel, ["vehiculo"])
        col_cedula_cliente = _find_col(sel, ["Cedula", "Cédula"]) or _find_col_contains(sel, ["cedula"])
        col_correo_cliente = _find_col(sel, ["correo", "Correo"]) or _find_col_contains(sel, ["correo"])

        cedula_cliente_value = _join_unique_values(sel[col_cedula_cliente].tolist()) if col_cedula_cliente else ""
        cliente_lookup = _lookup_cliente_info(ref_input, cedula_cliente_value)

        suma_comisiones_total = float(comision_exito) + float(plan_df["Comisión Mensual"].sum())

        template_context_default = _build_document_context(
            referencia=ref_input,
            bancos=sel[col_banco].astype(str).tolist(),
            pago_banco=pago_banco,
            comision_total=comision_exito,
            nombre_cliente=_join_unique_values(sel[col_nombre_cliente].tolist()) if col_nombre_cliente else "",
            numero_producto=_join_unique_values(sel[col_numero_producto].tolist()) if col_numero_producto else "",
            vehiculo=_join_unique_values(sel[col_vehiculo].tolist()) if col_vehiculo else "",
            cedula_cliente=cedula_cliente_value,
            correo_cliente=_join_unique_values(sel[col_correo_cliente].tolist()) if col_correo_cliente else "",
            telefono_cliente=cliente_lookup.get("telefono_cliente", ""),
            ciudad_cliente=cliente_lookup.get("ciudad_cliente", ""),
            direccion_cliente=cliente_lookup.get("direccion_cliente", ""),
            suma_comisiones_total=suma_comisiones_total,
        )
        template_context = _build_document_context_inputs(template_context_default)
        if st.session_state.get("modo_sin_portafolio", False) and cronogramas_individuales:
            deuda_ids = list(cronogramas_individuales.keys())
            deuda_principal = deuda_ids[0]
            cronograma_principal = cronogramas_individuales[deuda_principal]
            sel_principal = sel[sel[col_id].astype(str) == str(deuda_principal)].copy()
            numero_producto_principal = _join_unique_values(sel_principal[col_numero_producto].tolist()) if (col_numero_producto and not sel_principal.empty) else ""
            bancos_principal = sel_principal[col_banco].astype(str).tolist() if not sel_principal.empty else sel[col_banco].astype(str).tolist()
            context_principal = template_context.copy()
            context_principal["numero_producto"] = numero_producto_principal or context_principal.get("numero_producto", "")
            context_principal["pago_banco"] = _format_currency_cop(float(cronograma_principal.loc[cronograma_principal["Concepto"].str.contains("Entidad Financiera", na=False), "Cantidad"].sum()))
            context_principal["comision_total"] = _format_currency_cop(float(cronograma_principal.loc[cronograma_principal["Concepto"].str.contains("Comisión Resuelve", na=False), "Cantidad"].sum()))
            export_docx_bytes = build_recaudo_docx(
                template_path=DOCX_TEMPLATE_PATH,
                cronograma_df=cronograma_principal,
                plan_df=plan_df.drop(columns=["plan_key"], errors="ignore"),
                template_context={**context_principal, "entidad_financiera": _join_unique_values(bancos_principal)},
                include_graduation_section=bool(st.session_state.get("doc_graduacion_check", False) and st.session_state.get("doc_graduacion_confirmada", False)),
            )

            extra_blocks = []
            for deuda_id in deuda_ids[1:]:
                cronograma_deuda = cronogramas_individuales[deuda_id]
                sel_deuda = sel[sel[col_id].astype(str) == str(deuda_id)].copy()
                numero_producto_deuda = _join_unique_values(sel_deuda[col_numero_producto].tolist()) if (col_numero_producto and not sel_deuda.empty) else ""
                bancos_deuda = sel_deuda[col_banco].astype(str).tolist() if not sel_deuda.empty else sel[col_banco].astype(str).tolist()
                context_deuda = template_context.copy()
                context_deuda["numero_producto"] = numero_producto_deuda or context_deuda.get("numero_producto", "")
                context_deuda["pago_banco"] = _format_currency_cop(float(cronograma_deuda.loc[cronograma_deuda["Concepto"].str.contains("Entidad Financiera", na=False), "Cantidad"].sum()))
                context_deuda["comision_total"] = _format_currency_cop(float(cronograma_deuda.loc[cronograma_deuda["Concepto"].str.contains("Comisión Resuelve", na=False), "Cantidad"].sum()))
                deuda_docx = build_recaudo_docx(
                    template_path=DOCX_TEMPLATE_PATH,
                    cronograma_df=cronograma_deuda,
                    plan_df=construir_plan_liquidacion(cronograma_deuda, comision_mensual).drop(columns=["plan_key"], errors="ignore"),
                    template_context={**context_deuda, "entidad_financiera": _join_unique_values(bancos_deuda)},
                    include_graduation_section=bool(st.session_state.get("doc_graduacion_check", False) and st.session_state.get("doc_graduacion_confirmada", False)),
                )
                extra_blocks.append(_extract_first_liquidacion_block_elements(Document(BytesIO(deuda_docx))))
            export_docx_bytes = insert_extra_liquidacion_blocks(export_docx_bytes, extra_blocks)
        else:
            export_docx_bytes = build_recaudo_docx(
                template_path=DOCX_TEMPLATE_PATH,
                cronograma_df=cronograma_editado,
                plan_df=plan_df.drop(columns=["plan_key"], errors="ignore"),
                template_context=template_context,
                include_graduation_section=bool(st.session_state.get("doc_graduacion_check", False) and st.session_state.get("doc_graduacion_confirmada", False)),
            )
        export_pdf_bytes = convert_docx_bytes_to_pdf_bytes(export_docx_bytes)
    except Exception as export_exc:
        export_pdf_error = str(export_exc)
        st.error(f"No pude preparar el documento PDF: {export_exc}")

if export_pdf_bytes:
    missing_document_fields = _missing_document_fields(template_context)
    suma_comision_resuelve = float(
        cronograma_editado.loc[
            cronograma_editado["Concepto"].str.contains("Comisión Resuelve", na=False),
            "Cantidad",
        ].sum()
    )
    suma_pago_entidad = float(
        cronograma_editado.loc[
            cronograma_editado["Concepto"].str.contains("Entidad Financiera", na=False),
            "Cantidad",
        ].sum()
    )

    if suma_comision_resuelve > float(comision_exito) + 0.01:
        missing_document_fields.append("Ajustar cronograma: la suma de Comisión Resuelve no puede ser mayor a la Comisión de éxito total")
    if suma_pago_entidad > float(pago_banco) + 0.01:
        missing_document_fields.append("Ajustar cronograma: Pago a Entidad Financiera no puede ser mayor a PAGO BANCO")
    referencia_export = re.sub(r"[^A-Za-z0-9._-]+", " ", str(ref_input or "sin referencia")).strip() or "sin referencia"
    export_filename = f"{date.today().isoformat()} - ref {referencia_export}.pdf"
    if missing_document_fields:
        st.warning("Completa o corrige estos puntos antes de descargar el PDF: " + ", ".join(missing_document_fields))
    st.download_button(
        "⬇️ Descargar PDF con tablas",
        data=export_pdf_bytes,
        file_name=export_filename,
        mime="application/pdf",
        use_container_width=True,
        disabled=bool(missing_document_fields),
    )
else:
    if not export_pdf_error:
        st.info("Primero completa datos suficientes en el cronograma y en el plan para generar el PDF.")
#############################################################################################################################################################################
#############################################################################################################################################################################

#############################################################################################################################################################################    
# =================== 7) Predicción con el modelo ===================
st.markdown("### 7) Predicción de **recaudo_real** con MLP")

def _to_float_or_nan(x):
    try:
        return float(x)
    except Exception:
        return np.nan

feature_vals = {
    "PRI-ULT": _to_float_or_nan(plazo),
    "Ratio_PP": _to_float_or_nan(pct_primer_pago if not np.isnan(pct_primer_pago) else np.nan),
    "C/A": _to_float_or_nan(cuota_apartado if not np.isnan(cuota_apartado) else np.nan),
    "AMOUNT_TOTAL": _to_float_or_nan(comision_exito),
}

with st.expander("🔎 Ver features que se enviarán al modelo (crudas)", expanded=False):
    st.json(feature_vals)

issues = []
if np.isnan(feature_vals["AMOUNT_TOTAL"]) or feature_vals["AMOUNT_TOTAL"] < 0:
    issues.append("AMOUNT_TOTAL (Comisión de éxito total) no puede ser NaN ni negativa.")
if np.isnan(feature_vals["PRI-ULT"]) or feature_vals["PRI-ULT"] < 1:
    issues.append("PRI-ULT (PLAZO) debe ser un entero ≥ 1.")
if np.isnan(feature_vals["Ratio_PP"]) or feature_vals["Ratio_PP"] < 0:
    issues.append("Ratio_PP (% Primer Pago) no puede ser NaN ni negativo. Usa 0 si aplica.")
if np.isnan(feature_vals["C/A"]) or feature_vals["C/A"] <= 0:
    issues.append("C/A (Cuota/Apartado) debe ser > 0.")

if issues:
    st.warning("⚠️ Revisa antes de predecir:\n- " + "\n- ".join(issues))

with st.expander("📊 Diagnóstico de guardado en Google Sheets", expanded=False):
    gs_status = google_sheets_status()

    if gs_status["ok"]:
        st.success(
            "Conexión lista. "
            f"Cuenta de servicio: `{gs_status['client_email']}` • "
            f"Archivo: `{gs_status['spreadsheet_title']}` • "
            f"Hoja: `{gs_status['worksheet_title']}`"
        )
        st.caption(
            "No necesitas escribir los encabezados manualmente: la app intenta ponerlos en la fila 1 "
            "cuando guardas el primer registro."
        )
    else:
        st.error(
            "La app no pudo conectarse a Google Sheets. "
            f"Detalle: {gs_status['error']}"
        )
        st.markdown(
            f"""
**Revisa estos puntos:**
- El secreto debe existir en **Streamlit Secrets del despliegue actual** (sandbox/prod tienen secretos separados).
- Puedes usar `MI_JSON` (recomendado) o `GOOGLE_SERVICE_ACCOUNT_JSON`.
- Debes compartir el spreadsheet con este correo: `{gs_status['client_email']}`.
- La pestaña debe llamarse exactamente `{GOOGLE_SHEET_TAB}`.
- El guardado solo se ejecuta cuando presionas **Predecir recaudo**.
"""
        )

col_pred1, col_pred2 = st.columns([1,1])
with col_pred1:
    do_predict = st.button("🔮 Predecir recaudo", type="primary", use_container_width=True)
with col_pred2:
    meta = load_meta(_meta_version())
    if meta:
        st.caption(f"Modelo cargado • target: {meta.get('target','recaudo_real')}")

if do_predict:
    try:
        model = load_model(_model_version())
        FEATURES_RAW = ["PRI-ULT", "Ratio_PP", "C/A", "AMOUNT_TOTAL"]
        X_pred = pd.DataFrame([feature_vals], columns=FEATURES_RAW)
        yhat = float(model.predict(X_pred)[0])

# Ajustes existentes
        if yhat == 0.98:
            yhat_adj = yhat + 0.02
        elif yhat == 0.99:
            yhat_adj = yhat + 0.01
        else:
            yhat_adj = yhat + 0.03

# ✅ NUEVO AJUSTE
        if feature_vals["AMOUNT_TOTAL"] > 6_000_000:
            yhat_adj += 0.05

        st.success(f"✅ Predicción de recaudo: {yhat_adj:,.2f}")

        # ✅ Guardar registro del cálculo
        log_result = guardar_log_calculo(
            referencia=ref_input,
            ids=ids_sel,
            features=feature_vals,
            prediccion=yhat_adj,
        )
        st.session_state.last_prediction_value = float(yhat_adj)
        st.session_state.last_prediction_ready = True

        if log_result["sheet_ok"]:
            st.caption(f"📊 Histórico guardado en: `{log_result['sheet_destination']}`")
        else:
            st.warning(
                "No se pudo guardar el histórico en Google Sheets. "
                f"Detalle: {log_result['sheet_error']}"
            )

        if log_result["local_ok"]:
            st.caption(f"🗂️ Respaldo local guardado en: `{log_result['local_path']}`")
        elif not log_result["sheet_ok"]:
            st.error(
                "Tampoco se pudo guardar el respaldo local. "
                f"Detalle: {log_result['local_error']}"
            )

        st.caption("Entradas usadas por el pipeline (crudas):")
        st.dataframe(pd.DataFrame([feature_vals]), use_container_width=True)

    except Exception as e:
        st.error(f"Error al predecir: {e}")

st.markdown("---")
st.markdown("### 8) Envío a aprobación de estructurados")
st.caption("Este envío se hace solo cuando presionas el botón de aprobación.")

st.markdown("#### Adjuntos obligatorios (Drive con autenticación de usuario)")
st.caption(
    "Flujo automático: autentica tu cuenta Google y sube los adjuntos obligatorios."
)

oauth_disponible = _oauth_drive_configurado()
carta_pagare_file = None
pantallazo_file = None
condonacion_correo_file = None
if not oauth_disponible:
    st.error(
        "Este despliegue no tiene secretos OAuth configurados. "
        "Configura GOOGLE_OAUTH_CLIENT_ID y GOOGLE_OAUTH_CLIENT_SECRET."
    )
else:
    auth_col1, auth_col2 = st.columns([1, 2])
    with auth_col1:
        iniciar_auth_drive = st.button("🔐 Iniciar autenticación Drive", use_container_width=True)
    with auth_col2:
        if st.session_state.get("drive_user_token"):
            st.success("Cuenta Drive autenticada en esta sesión.")
        else:
            st.info("Sin autenticación activa.")

    if iniciar_auth_drive:
        try:
            flow, auth_url, oauth_state, client_config = _get_drive_flow_and_auth_url()
            st.session_state.drive_oauth_state = oauth_state
            st.session_state.drive_oauth_code_verifier = getattr(flow, "code_verifier", None)
            st.session_state.drive_auth_url = auth_url
            st.session_state.drive_auth_client_config = client_config
            st.session_state.drive_auth_in_progress = True
        except Exception as e:
            st.error(f"No se pudo iniciar autenticación OAuth: {e}")

    if st.session_state.get("drive_auth_in_progress"):
        auth_url = st.session_state.get("drive_auth_url", "")
        redirect_uri = _get_oauth_redirect_uri()
        query_code = str(st.query_params.get("code", "")).strip() if hasattr(st, "query_params") else ""
        last_processed_code = str(st.session_state.get("drive_last_processed_code", "")).strip()

        if auth_url:
            if "localhost" in redirect_uri:
                st.markdown(
                    f"1) Abre este enlace y autoriza tu cuenta: [Autorizar Drive]({auth_url})  \n"
                    "2) Copia el `code` de la URL de redirección (o pega la URL completa)."
                )
            else:
                st.markdown(
                    f"1) Abre este enlace y autoriza tu cuenta: [Autorizar Drive]({auth_url})  \n"
                    "2) Serás redirigido automáticamente a esta app."
                )

        if query_code and "localhost" not in redirect_uri and query_code != last_processed_code:
            try:
                _complete_drive_oauth_with_code(query_code)
                st.session_state.drive_last_processed_code = query_code
                st.success("Autenticación de Drive completada automáticamente.")
                try:
                    st.query_params.clear()
                except Exception:
                    pass
            except Exception as e:
                st.error(f"No fue posible completar OAuth automáticamente: {e}")

        auth_code_input = st.text_input(
            "Código OAuth / URL de redirección",
            value=query_code,
            key="drive_oauth_code_input",
            help="Si usas redirect localhost, pega la URL completa o solo el code.",
        )
        finalizar_auth_drive = st.button("✅ Confirmar autenticación", use_container_width=True)
        if finalizar_auth_drive:
            try:
                code = _extract_oauth_code(auth_code_input)
                if not code:
                    st.warning("Debes pegar un código OAuth válido.")
                else:
                    _complete_drive_oauth_with_code(code)
                    st.session_state.drive_last_processed_code = code
                    st.success("Autenticación de Drive completada.")
            except Exception as e:
                st.error(f"No fue posible completar OAuth: {e}")

    carta_pagare_file = st.file_uploader(
        "📎 Adjuntar carta con pagaré firmado (PDF)",
        type=["pdf"],
        key="carta_pagare_pdf",
    )
    pantallazo_file = st.file_uploader(
        "📎 Adjuntar pantallazo de aceptación del cliente (PDF o imagen)",
        type=["pdf", "png", "jpg", "jpeg", "webp"],
        key="pantallazo_pdf",
    )
    
correo_para_sheets = st.text_input(
    "📧 Dirección de correo electrónico (obligatorio para enviar)",
    key="correo_para_sheets",
).strip()
condonacion_mensualidades = st.selectbox(
    "¿El cliente cuenta con condonación de mensualidades? (obligatorio)",
    options=["", "Si", "No"],
    index=0,
    key="condonacion_mensualidades",
)
if condonacion_mensualidades == "Si":
    condonacion_correo_file = st.file_uploader(
        "📎 Adjuntar pantallazo de correo de aprobación de condonación (PDF o imagen)",
        type=["pdf", "png", "jpg", "jpeg", "webp"],
        key="condonacion_correo_soporte",
    )

enviar_aprobacion = st.button("Enviar AProbación estructurados", use_container_width=True)
if enviar_aprobacion:
    pred_value = st.session_state.get("last_prediction_value")
    carta_link_final = ""
    pantallazo_link_final = ""
    condonacion_correo_link_final = ""
    duplicate_mode = "none"
    duplicate_exact_rows = []
    duplicate_key = f"{datetime.now().strftime('%Y-%m')}|{_norm(ref_input)}|{'-'.join(sorted(_ids_to_set(ids_sel)))}"
    if pred_value is None:
        st.warning("Primero debes presionar **Predecir recaudo**.")
    elif not correo_para_sheets:
        st.warning("Debes ingresar el correo electrónico antes de enviar.")
    elif condonacion_mensualidades not in {"Si", "No"}:
        st.warning("Debes seleccionar Si o No en condonación de mensualidades.")
    elif carta_pagare_file is None or pantallazo_file is None:
        st.warning("Debes adjuntar Carta/Pagaré (solo PDF) y Pantallazo de aceptación (PDF o imagen).")
    elif condonacion_mensualidades == "Si" and condonacion_correo_file is None:
        st.warning("Debes adjuntar el pantallazo de correo de aprobación de condonación (PDF o imagen).")
    else:
        pred_value_envio = round(float(pred_value or 0.0), 4)
        st.info(
            f"Predicción de recaudo usada para este envío: **{pred_value_envio:.4f}** "
            "(este mismo valor se guarda en la última columna)."
        )
        duplicate_check = _get_respuestas_duplicados_mes(ref_input, ids_sel)
        if not duplicate_check["ok"]:
            st.error(
                "No fue posible validar duplicados contra la hoja de respuestas. "
                f"Detalle: {duplicate_check['error']}"
            )
            st.stop()

        duplicate_mode = duplicate_check["mode"]
        duplicate_exact_rows = duplicate_check.get("exact_rows", [])
        confirmed_key = str(st.session_state.get("duplicate_confirm_key", ""))
        if duplicate_mode == "exact_duplicate" and confirmed_key != duplicate_key:
            st.session_state.duplicate_confirm_key = duplicate_key
            st.warning(
                "⚠️ Esta referencia con el/los mismo(s) ID(s) ya fue subida este mes. "
                "Si deseas reenviarla a aprobación, presiona nuevamente **Enviar AProbación estructurados**."
            )
            st.stop()
        if duplicate_mode == "reference_duplicate":
            st.warning(
                "ℹ️ Esta referencia ya fue originada este mes con otro ID de deuda. "
                "Se enviará sin check de aprobación estructurados y sin comentario. "
                "Por favor contacta al equipo de estructurados."
            )
            
        is_valid_pdf, pdf_validation_message = _validate_carta_pagare_pdf(
            carta_pagare_file,
            ref_input,
            expected_flow_for_pdf_validation,
        )
        if not is_valid_pdf:
            st.warning(pdf_validation_message)
            st.stop()
        try:
            drive_service = _build_drive_service_from_session()
            if drive_service is None:
                st.warning("Debes autenticar Drive antes de enviar la aprobación.")
                st.stop()

            carta_upload = _upload_file_to_drive(
                drive_service,
                carta_pagare_file,
                DRIVE_FOLDER_CARTA_PAGARE_ID,
                allowed_extensions=(".pdf",),
                invalid_message="Carta con pagaré firmado: solo se permite PDF.",
            )
            pantallazo_upload = _upload_file_to_drive(
                drive_service,
                pantallazo_file,
                DRIVE_FOLDER_PANTALLAZOS_ID,
                allowed_extensions=(".pdf", ".png", ".jpg", ".jpeg", ".webp"),
                invalid_message="Pantallazo de aceptación: solo PDF o imagen (PNG/JPG/JPEG/WEBP).",
            )
            carta_link_final = carta_upload.get("webViewLink", "")
            pantallazo_link_final = pantallazo_upload.get("webViewLink", "")
            st.caption(f"📂 Carta/Pagaré cargado: {carta_link_final}")
            st.caption(f"📂 Pantallazo cargado: {pantallazo_link_final}")
            if condonacion_mensualidades == "Si" and condonacion_correo_file is not None:
                condonacion_upload = _upload_file_to_drive(
                    drive_service,
                    condonacion_correo_file,
                    DRIVE_FOLDER_CONDONACION_CORREO_ID,
                    allowed_extensions=(".pdf", ".png", ".jpg", ".jpeg", ".webp"),
                    invalid_message="Pantallazo de correo de condonación: solo PDF o imagen (PNG/JPG/JPEG/WEBP).",
                )
                condonacion_correo_link_final = condonacion_upload.get("webViewLink", "")
                st.caption(f"📂 Correo aprobación condonación cargado: {condonacion_correo_link_final}")
            elif condonacion_mensualidades == "No":
                condonacion_correo_link_final = ""
        except Exception as e:
            st.error(f"No se pudieron subir los adjuntos a Drive: {e}")
            st.stop()

    if (
        pred_value is not None
        and correo_para_sheets
        and condonacion_mensualidades in {"Si", "No"}
        and carta_link_final
        and pantallazo_link_final
        and (condonacion_mensualidades == "No" or bool(condonacion_correo_link_final))
    ):
            
        envio_result = enviar_aprobacion_estructurados(
            referencia=ref_input,
            ids=ids_sel,
            bancos=bancos_sel_text,
            correo_electronico=correo_para_sheets,
            condonacion_mensualidades=condonacion_mensualidades,
            comision_exito_total=feature_vals.get("AMOUNT_TOTAL"),
            ce_inicial=ce_inicial,
            prediccion=pred_value_envio,
            tipo_liquidacion=tipo_liquidacion_val,
            carta_pagare_link=carta_link_final,
            pantallazo_aceptacion_link=pantallazo_link_final,
            pantallazo_correo_condonacion_link=condonacion_correo_link_final,
            duplicate_mode=duplicate_mode,
            exact_rows_previas=duplicate_exact_rows,
        )
        if envio_result["estr_ok"]:
            st.session_state.duplicate_confirm_key = ""
            estado_aprob = "✅ Aprobado" if envio_result["es_aprobado"] else "⛔ No aprobado"
            st.success(f"Envío exitoso a `{envio_result['estr_destination']}`.")
            if envio_result.get("duplicate_mode") == "exact_duplicate":
                st.info(
                    "Se detectó duplicado exacto: se desmarcó la aprobación previa "
                    "y el comentario anterior aprobado quedó como 'Duplicado'."
                )
            elif envio_result.get("duplicate_mode") == "reference_duplicate":
                st.info("Registro enviado sin aprobación estructurados ni comentario por referencia repetida del mes.")
            st.caption(
                f"Criterio aplicado ({tipo_liquidacion_val or 'No Tradicional'}): "
                f"predicción {pred_value_envio:.2f} vs umbral {envio_result['umbral_aprobacion']:.2f} → {estado_aprob}"
            )
        else:
            st.error(
                "No se pudo enviar a estructurados. "
                f"Detalle: {envio_result['estr_error']}"
            )
